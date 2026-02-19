"""Tests for .docx document generation — comparison, payment schedule, company statement."""

from __future__ import annotations

import os
import re
from datetime import date
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import RGBColor

from services.comparison_engine import ComparisonEngine
from services.document_generator import DocumentGenerator, OUTPUT_DIR
from services.payment_schedule import PaymentScheduleGenerator

# ---------------------------------------------------------------------------
#  Test data — PBM reference (same as test_comparison_engine.py)
# ---------------------------------------------------------------------------

PBM_COMPANY_NAME = "Point Blank Medical Pty Ltd"
PBM_ACN = "123456789"
PBM_CREDITORS_TOTAL = 985_777.37

PBM_PLAN = {
    "total_contribution": 516_000.0,
    "practitioner_fee_pct": 10.0,
    "num_initial_payments": 2,
    "initial_payment_amount": 32_500.0,
    "num_ongoing_payments": 22,
    "ongoing_payment_amount": 20_500.0,
    "est_liquidator_fees": 50_000.0,
    "est_legal_fees": 10_000.0,
    "est_disbursements": 5_000.0,
}

PBM_ASSETS = [
    {
        "asset_type": "cash",
        "description": "Cash at Bank",
        "book_value": 59_689.27,
        "liquidation_recovery_pct": 0.20,
        "liquidation_value": 59_689.27 * 0.20,
    },
    {
        "asset_type": "receivables",
        "description": "Accounts Receivable",
        "book_value": 69_553.24,
        "liquidation_recovery_pct": 0.30,
        "liquidation_value": 69_553.24 * 0.30,
    },
    {
        "asset_type": "inventory",
        "description": "Inventory",
        "book_value": 51_826.62,
        "liquidation_recovery_pct": 0.25,
        "liquidation_value": 51_826.62 * 0.25,
    },
    {
        "asset_type": "loans_related",
        "description": "Loans to Related Entities",
        "book_value": 34_964.83,
        "liquidation_recovery_pct": 0.30,
        "liquidation_value": 34_964.83 * 0.30,
    },
    {
        "asset_type": "loans_shareholder",
        "description": "Shareholder Loans",
        "book_value": 2_010_000.00,
        "liquidation_recovery_pct": 0.00,
        "liquidation_value": 0.0,
    },
    {
        "asset_type": "equipment",
        "description": "Plant & Equipment",
        "book_value": 15_000.00,
        "liquidation_recovery_pct": 0.25,
        "liquidation_value": 15_000.00 * 0.25,
    },
]

# ---------------------------------------------------------------------------
#  Shared generators
# ---------------------------------------------------------------------------

comparison_engine = ComparisonEngine()
payment_generator = PaymentScheduleGenerator()
doc_generator = DocumentGenerator()


def _comparison_data() -> dict:
    return comparison_engine.calculate(PBM_ASSETS, PBM_CREDITORS_TOTAL, PBM_PLAN)


def _schedule_data() -> dict:
    return payment_generator.generate(PBM_PLAN)


def _sample_sections() -> list[dict]:
    """Build sample narrative sections for the company offer statement."""
    return [
        {
            "section": "background",
            "content": (
                "Point Blank Medical Pty Ltd was incorporated in 2018. "
                "The company operates in the medical supply sector. "
                "[REQUIRES INPUT: date of incorporation]"
            ),
            "status": "approved",
        },
        {
            "section": "expert_advice",
            "content": (
                "The directors sought professional advice in January 2026. "
                "A restructuring practitioner was appointed under Part 5.3B. "
                "[UNKNOWN TERM: Section 453B notice]"
            ),
            "status": "draft",
        },
        {
            "section": "plan_summary",
            "content": (
                "The proposed plan involves a total contribution of $516,000. "
                "The estimated dividend is 47.1 cents in the dollar."
            ),
            "status": "reviewed",
        },
        {
            "section": "viability",
            "content": (
                "The company has implemented cost reductions and secured "
                "new revenue streams. Operations are projected to be profitable."
            ),
            "status": "approved",
        },
        {
            "section": "comparison_commentary",
            "content": (
                "Under the SBR plan, creditors receive 47.1 cents in the dollar "
                "compared to nil in a liquidation scenario."
            ),
            "status": "draft",
        },
        {
            "section": "distress_events",
            "content": (
                "In mid-2025, the company lost a major supply contract, "
                "leading to cash flow difficulties. "
                "[REQUIRES INPUT: specific date of contract loss]"
            ),
            "status": "draft",
        },
    ]


# ---------------------------------------------------------------------------
#  Helper: extract all text from a docx
# ---------------------------------------------------------------------------

def _all_text(doc: DocxDocument) -> str:
    """Extract all paragraph and table text from a docx document."""
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


# ===========================================================================
#  Comparison Document Tests (1-6)
# ===========================================================================


class TestComparisonDocxGenerates:
    """1. test_comparison_docx_generates — file exists, is valid .docx."""

    def test_comparison_docx_generates(self):
        data = _comparison_data()
        filepath = doc_generator.generate_comparison_docx(
            data, PBM_COMPANY_NAME, acn=PBM_ACN
        )
        assert filepath.exists()
        assert filepath.suffix == ".docx"
        # Verify it can be opened as a valid docx
        doc = DocxDocument(str(filepath))
        assert len(doc.paragraphs) > 0
        # Cleanup
        os.remove(str(filepath))


class TestComparisonDocxHasTable:
    """2. test_comparison_docx_has_table — document contains a table."""

    def test_comparison_docx_has_table(self):
        data = _comparison_data()
        filepath = doc_generator.generate_comparison_docx(
            data, PBM_COMPANY_NAME, acn=PBM_ACN
        )
        doc = DocxDocument(str(filepath))
        assert len(doc.tables) >= 1
        os.remove(str(filepath))


class TestComparisonDocxPBMValues:
    """3. test_comparison_docx_pbm_values — table contains 47.1, 464400, 985777.37."""

    def test_comparison_docx_pbm_values(self):
        data = _comparison_data()
        filepath = doc_generator.generate_comparison_docx(
            data, PBM_COMPANY_NAME, acn=PBM_ACN
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        assert "47.1" in text
        assert "464,400.00" in text
        assert "985,777.37" in text
        os.remove(str(filepath))


class TestComparisonDocxHasNotes:
    """4. test_comparison_docx_has_notes — numbered notes section present."""

    def test_comparison_docx_has_notes(self):
        data = _comparison_data()
        filepath = doc_generator.generate_comparison_docx(
            data, PBM_COMPANY_NAME, acn=PBM_ACN
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        # Check for numbered notes
        assert "Note 1:" in text
        assert "Note 2:" in text
        assert "Notes" in text
        os.remove(str(filepath))


class TestComparisonDocxCurrencyFormat:
    """5. test_comparison_docx_currency_format — values formatted as $XXX,XXX.XX."""

    def test_comparison_docx_currency_format(self):
        data = _comparison_data()
        filepath = doc_generator.generate_comparison_docx(
            data, PBM_COMPANY_NAME, acn=PBM_ACN
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        # Check Australian currency format with commas
        assert re.search(r"\$\d{1,3}(,\d{3})*\.\d{2}", text)
        os.remove(str(filepath))


class TestComparisonDocxNegativeInParens:
    """6. test_comparison_docx_negative_in_parens — negative values in parentheses."""

    def test_comparison_docx_negative_in_parens(self):
        data = _comparison_data()
        filepath = doc_generator.generate_comparison_docx(
            data, PBM_COMPANY_NAME, acn=PBM_ACN
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        # Negative values should be in parens like ($50,000.00)
        assert re.search(r"\(\$[\d,]+\.\d{2}\)", text)
        # Should NOT have a minus-dollar pattern like -$50,000
        assert not re.search(r"-\$[\d,]+\.\d{2}", text)
        os.remove(str(filepath))


# ===========================================================================
#  Payment Schedule Tests (7-10)
# ===========================================================================


class TestPaymentScheduleDocxGenerates:
    """7. test_payment_schedule_docx_generates — file exists, is valid .docx."""

    def test_payment_schedule_docx_generates(self):
        data = _schedule_data()
        filepath = doc_generator.generate_payment_schedule_docx(
            data, PBM_COMPANY_NAME
        )
        assert filepath.exists()
        assert filepath.suffix == ".docx"
        doc = DocxDocument(str(filepath))
        assert len(doc.paragraphs) > 0
        os.remove(str(filepath))


class TestPaymentScheduleDocx24Rows:
    """8. test_payment_schedule_docx_24_rows — table has 24 data rows + header + total."""

    def test_payment_schedule_docx_24_rows(self):
        data = _schedule_data()
        filepath = doc_generator.generate_payment_schedule_docx(
            data, PBM_COMPANY_NAME
        )
        doc = DocxDocument(str(filepath))
        table = doc.tables[0]

        # 1 header row + 24 data rows + 1 total row = 26
        assert len(table.rows) == 26
        os.remove(str(filepath))


class TestPaymentScheduleDocxTotals:
    """9. test_payment_schedule_docx_totals — total row sums correctly to $516,000."""

    def test_payment_schedule_docx_totals(self):
        data = _schedule_data()
        filepath = doc_generator.generate_payment_schedule_docx(
            data, PBM_COMPANY_NAME
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        # Total contribution = $516,000.00
        assert "$516,000.00" in text
        # Total fees = $51,600.00
        assert "$51,600.00" in text
        # Total net dividend = $464,400.00
        assert "$464,400.00" in text
        os.remove(str(filepath))


class TestPaymentScheduleDocxInitialPayments:
    """10. test_payment_schedule_docx_initial_payments — first 2 rows show $32,500."""

    def test_payment_schedule_docx_initial_payments(self):
        data = _schedule_data()
        filepath = doc_generator.generate_payment_schedule_docx(
            data, PBM_COMPANY_NAME
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        assert "$32,500.00" in text
        os.remove(str(filepath))


# ===========================================================================
#  Company Statement Tests (11-15)
# ===========================================================================


class TestCompanyStatementDocxGenerates:
    """11. test_company_statement_docx_generates — file exists, is valid .docx."""

    def test_company_statement_docx_generates(self):
        sections = _sample_sections()
        filepath = doc_generator.generate_company_statement_docx(
            sections,
            PBM_COMPANY_NAME,
            acn=PBM_ACN,
            practitioner_name="Test Practitioner",
        )
        assert filepath.exists()
        assert filepath.suffix == ".docx"
        doc = DocxDocument(str(filepath))
        assert len(doc.paragraphs) > 0
        os.remove(str(filepath))


class TestCompanyStatement6Sections:
    """12. test_company_statement_6_sections — all 6 section headings present."""

    def test_company_statement_6_sections(self):
        sections = _sample_sections()
        filepath = doc_generator.generate_company_statement_docx(
            sections,
            PBM_COMPANY_NAME,
            acn=PBM_ACN,
            practitioner_name="Test Practitioner",
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        assert "SECTION I" in text
        assert "SECTION II" in text
        assert "SECTION III" in text
        assert "SECTION IV" in text
        assert "SECTION V" in text
        assert "SECTION VI" in text

        assert "BACKGROUND" in text
        assert "EXPERT ADVICE AND APPOINTMENT" in text
        assert "THE RESTRUCTURING PLAN" in text
        assert "VIABILITY AND FUTURE OPERATIONS" in text
        assert "COMPARISON OF OUTCOMES" in text
        assert "DISTRESS EVENTS" in text
        os.remove(str(filepath))


class TestCompanyStatementDraftWatermark:
    """13. test_company_statement_draft_watermark — unapproved sections marked as DRAFT."""

    def test_company_statement_draft_watermark(self):
        sections = _sample_sections()
        filepath = doc_generator.generate_company_statement_docx(
            sections,
            PBM_COMPANY_NAME,
            acn=PBM_ACN,
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        # Sections with status != "approved" should have draft marker
        assert "[DRAFT" in text
        assert "NOT YET APPROVED" in text

        # Count how many draft markers there are
        # expert_advice (draft), plan_summary (reviewed), comparison_commentary (draft),
        # distress_events (draft) = 4 sections should have draft markers
        draft_count = text.count("[DRAFT")
        assert draft_count == 4
        os.remove(str(filepath))


class TestCompanyStatementRequiresInputHighlighted:
    """14. test_company_statement_requires_input_highlighted — [REQUIRES INPUT] flags highlighted."""

    def test_company_statement_requires_input_highlighted(self):
        sections = _sample_sections()
        filepath = doc_generator.generate_company_statement_docx(
            sections,
            PBM_COMPANY_NAME,
            acn=PBM_ACN,
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        # Should contain the REQUIRES INPUT flags
        assert "[REQUIRES INPUT:" in text

        # Check that the flag runs are bold and have red color
        found_highlighted = False
        for para in doc.paragraphs:
            for run in para.runs:
                if "[REQUIRES INPUT:" in run.text:
                    assert run.bold is True
                    assert run.font.color.rgb == RGBColor(0xB9, 0x1C, 0x1C)
                    found_highlighted = True

        assert found_highlighted, "No highlighted [REQUIRES INPUT] flags found"

        # Also check for UNKNOWN TERM highlighting
        found_unknown = False
        for para in doc.paragraphs:
            for run in para.runs:
                if "[UNKNOWN TERM:" in run.text:
                    assert run.bold is True
                    assert run.font.color.rgb == RGBColor(0xB9, 0x1C, 0x1C)
                    found_unknown = True

        assert found_unknown, "No highlighted [UNKNOWN TERM] flags found"
        os.remove(str(filepath))


class TestCompanyStatementRomanNumerals:
    """15. test_company_statement_roman_numerals — sections numbered I through VI."""

    def test_company_statement_roman_numerals(self):
        sections = _sample_sections()
        filepath = doc_generator.generate_company_statement_docx(
            sections,
            PBM_COMPANY_NAME,
            acn=PBM_ACN,
        )
        doc = DocxDocument(str(filepath))
        text = _all_text(doc)

        numerals = ["I", "II", "III", "IV", "V", "VI"]
        for numeral in numerals:
            assert f"SECTION {numeral} " in text or f"SECTION {numeral}\n" in text, \
                f"Roman numeral {numeral} not found"
        os.remove(str(filepath))


# ===========================================================================
#  Existing Generator Tests (16-17) — Rule 1 compliance
# ===========================================================================


class TestExistingDIRRIStillWorks:
    """16. test_existing_dirri_still_works — existing DIRRI generation unbroken."""

    def test_existing_dirri_still_works(self):
        from models.schemas import (
            AppointmentType,
            CompanyData,
            DIRRIRequest,
            FirmProfile,
        )

        firm = FirmProfile(
            firm_name="Test Firm",
            practitioner_name="Test Practitioner",
        )
        company = CompanyData(
            legal_name="Test Company Pty Ltd",
            acn="111222333",
        )
        request = DIRRIRequest(
            firm_profile=firm,
            company=company,
            appointment_type=AppointmentType.SMALL_BUSINESS_RESTRUCTURING,
            appointment_date=date(2026, 1, 15),
        )

        filepath = doc_generator.generate_dirri(request)
        assert filepath.exists()
        assert filepath.suffix == ".docx"

        doc = DocxDocument(str(filepath))
        text = _all_text(doc)
        assert "DIRRI" in text or "Declaration of Independence" in text
        assert "Test Company Pty Ltd" in text
        os.remove(str(filepath))


class TestExistingSafeHarbourStillWorks:
    """17. test_existing_safe_harbour_still_works — existing Safe Harbour generation unbroken."""

    def test_existing_safe_harbour_still_works(self):
        from models.schemas import CompanyData, FirmProfile

        firm = FirmProfile(
            firm_name="Test Firm",
            practitioner_name="Test Practitioner",
        )
        company = CompanyData(
            legal_name="Test Company Pty Ltd",
            acn="111222333",
        )

        filepath = doc_generator.generate_safe_harbour_checklist(firm, company)
        assert filepath.exists()
        assert filepath.suffix == ".docx"

        doc = DocxDocument(str(filepath))
        text = _all_text(doc)
        assert "Safe Harbour" in text
        assert "Test Company Pty Ltd" in text
        os.remove(str(filepath))
