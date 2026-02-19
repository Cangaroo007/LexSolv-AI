"""
PBM Regression Test Suite
========================
These tests validate the entire LexSolv pipeline against Point Blank Medical
reference data. If any test in this file fails after a code change, the change
broke something critical.

Golden values (from real PBM case):
- Total creditor claims: $985,777.37
- SBR contribution: $516,000
- Practitioner fee: 10%
- SBR available for distribution: $464,400
- SBR dividend: 47.1 cents in the dollar
- Liquidation dividend: 0 cents
- Initial payments: 2 x $32,500
- Ongoing payments: 22 x $20,500
"""

from __future__ import annotations

import os
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from services.file_parser import FileParser
from services.creditor_schedule import CreditorScheduleService
from services.comparison_engine import ComparisonEngine
from services.payment_schedule import PaymentScheduleGenerator
from services.document_generator import DocumentGenerator
from services.privacy_vault import scrub, restore

# ---------------------------------------------------------------------------
#  Fixtures & Constants
# ---------------------------------------------------------------------------

FIXTURES = Path(__file__).parent / "fixtures"

PBM_AGED_PAYABLES = str(FIXTURES / "pbm_aged_payables.csv")
PBM_BALANCE_SHEET = str(FIXTURES / "pbm_balance_sheet.csv")
PBM_DIRECTOR_NARRATIVE = str(FIXTURES / "pbm_director_narrative.txt")

PBM_COMPANY_NAME = "Point Blank Medical Pty Ltd"
PBM_ACN = "123456789"
PBM_CREDITORS_TOTAL = 985_777.37

PBM_PLAN = {
    "total_contribution": 516_000.0,
    "practitioner_fee_pct": 10.0,
    "num_initial_payments": 2,
    "initial_payment_amount": 32_500.0,
    "num_ongoing_payments": 22,
    "ongoing_payment_amount": 20_500.0,
    "est_liquidator_fees": 50_000.0,
    "est_legal_fees": 10_000.0,
    "est_disbursements": 5_000.0,
}

# Expected creditor data from aged payables CSV
EXPECTED_CREDITORS = {
    "Australian Taxation Office - ITA": 573_230.31,
    "Australian Taxation Office - ICA": 268_294.01,
    "iCare NSW": 825.23,
    "Prospa Advance": 143_874.02,
    "BlueShak": 142_105.81,
    "BTC Health Australia": 67_447.99,
}

# Expected balance sheet values
EXPECTED_ASSETS = {
    "cash": 59_689.27,
    "receivables": 69_553.24,
    "inventory": 51_826.62,
    "loans_to_related": 34_964.83,
    "loans_shareholder": 2_010_000.00,
    "equipment": 15_000.00,
}

# Expected liquidation recovery at default rates
EXPECTED_LIQUIDATION_RECOVERIES = {
    "cash": 59_689.27 * 0.20,        # $11,937.854
    "receivables": 69_553.24 * 0.30,  # $20,865.972
    "inventory": 51_826.62 * 0.25,    # $12,956.655
    "equipment": 15_000.00 * 0.25,    # $3,750.00
    "loans_related": 34_964.83 * 0.30,  # $10,489.449
    "loans_shareholder": 2_010_000.00 * 0.00,  # $0.00
}

# Service singletons
parser = FileParser()
creditor_svc = CreditorScheduleService()
comparison_engine = ComparisonEngine()
payment_generator = PaymentScheduleGenerator()
doc_generator = DocumentGenerator()


# ---------------------------------------------------------------------------
#  Helper: extract all text from a docx
# ---------------------------------------------------------------------------

def _all_text(doc: DocxDocument) -> str:
    """Extract all paragraph and table text from a docx document."""
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


def _load_narrative() -> str:
    """Load the PBM director narrative fixture."""
    return Path(PBM_DIRECTOR_NARRATIVE).read_text()


# ===========================================================================
#  FILE PARSING REGRESSION (Tests 1-4)
# ===========================================================================


class TestPBMAgedPayablesParse:
    """Tests 1-2: Parse aged payables CSV and verify all creditors extracted."""

    def test_pbm_aged_payables_parse_complete(self):
        """1. Parse pbm_aged_payables.csv -> verify ALL 6 creditors extracted."""
        result = parser.parse_aged_payables(PBM_AGED_PAYABLES)
        assert len(result) == 6
        names = {c["creditor_name"] for c in result}
        assert names == set(EXPECTED_CREDITORS.keys())

    def test_pbm_aged_payables_amounts_exact(self):
        """2. Verify each creditor amount matches to the cent."""
        result = parser.parse_aged_payables(PBM_AGED_PAYABLES)
        by_name = {c["creditor_name"]: c["amount_claimed"] for c in result}
        for name, expected_amount in EXPECTED_CREDITORS.items():
            assert by_name[name] == pytest.approx(expected_amount, abs=0.01), \
                f"{name}: expected {expected_amount}, got {by_name[name]}"


class TestPBMBalanceSheetParse:
    """Tests 3-4: Parse balance sheet CSV and verify all asset categories."""

    def test_pbm_balance_sheet_parse_complete(self):
        """3. Parse pbm_balance_sheet.csv -> verify ALL asset categories extracted."""
        result = parser.parse_balance_sheet(PBM_BALANCE_SHEET)
        for key in EXPECTED_ASSETS:
            assert key in result, f"Missing asset category: {key}"
            assert result[key] > 0, f"Asset {key} should be > 0"

    def test_pbm_balance_sheet_amounts_exact(self):
        """4. Verify each asset book value matches to the cent."""
        result = parser.parse_balance_sheet(PBM_BALANCE_SHEET)
        for key, expected_val in EXPECTED_ASSETS.items():
            assert result[key] == pytest.approx(expected_val, abs=0.01), \
                f"{key}: expected {expected_val}, got {result[key]}"


# ===========================================================================
#  CREDITOR SCHEDULE REGRESSION (Tests 5-8)
# ===========================================================================


class TestPBMCreditorClassification:
    """Tests 5-8: Auto-classify PBM creditors and verify totals."""

    def test_pbm_creditor_classification(self):
        """5. Auto-classify PBM creditors into correct categories."""
        parsed = parser.parse_aged_payables(PBM_AGED_PAYABLES)
        creditors = creditor_svc.build_from_parsed(parsed)
        by_name = {c["creditor_name"]: c for c in creditors}

        assert by_name["Australian Taxation Office - ITA"]["category"] == "ato_ita"
        assert by_name["Australian Taxation Office - ICA"]["category"] == "ato_ica"
        assert by_name["iCare NSW"]["category"] == "workers_comp"
        assert by_name["Prospa Advance"]["category"] == "finance"
        assert by_name["BlueShak"]["category"] == "trade"
        assert by_name["BTC Health Australia"]["category"] == "trade"

    def test_pbm_related_party_flagging(self):
        """6. BlueShak flagged as related party -> can_vote=False."""
        parsed = parser.parse_aged_payables(PBM_AGED_PAYABLES)
        creditors = creditor_svc.build_from_parsed(parsed)
        by_name = {c["creditor_name"]: c for c in creditors}

        blueshak = by_name["BlueShak"]
        creditor_svc.flag_related_party(blueshak, True)
        assert blueshak["is_related_party"] is True
        assert blueshak["can_vote"] is False

    def test_pbm_debt_forgiveness(self):
        """7. BlueShak marked as forgiven -> excluded from admitted totals."""
        parsed = parser.parse_aged_payables(PBM_AGED_PAYABLES)
        creditors = creditor_svc.build_from_parsed(parsed)
        by_name = {c["creditor_name"]: c for c in creditors}

        blueshak = by_name["BlueShak"]
        creditor_svc.flag_related_party(blueshak, True)
        creditor_svc.update_status(blueshak, "forgiven")

        totals = creditor_svc.calculate_totals(creditors)
        assert totals["total_excluded"] == 1
        # BlueShak's amount should NOT be in total_admitted
        assert totals["total_admitted"] == pytest.approx(
            sum(EXPECTED_CREDITORS.values()) - EXPECTED_CREDITORS["BlueShak"],
            abs=0.01,
        )

    def test_pbm_total_claims(self):
        """8. Total claims from balance sheet = $985,777.37."""
        result = parser.parse_balance_sheet(PBM_BALANCE_SHEET)
        assert result["total_liabilities"] == pytest.approx(PBM_CREDITORS_TOTAL, abs=0.01)


# ===========================================================================
#  COMPARISON ENGINE REGRESSION (Tests 9-13)
# ===========================================================================


class TestPBMComparisonEngine:
    """Tests 9-13: SBR and liquidation dividend calculations."""

    def _build_pbm_assets(self) -> list[dict]:
        parsed_bs = parser.parse_balance_sheet(PBM_BALANCE_SHEET)
        return comparison_engine.build_assets_from_balance_sheet(parsed_bs)

    def _calculate(self) -> dict:
        assets = self._build_pbm_assets()
        return comparison_engine.calculate(assets, PBM_CREDITORS_TOTAL, PBM_PLAN)

    def test_pbm_sbr_dividend_exact(self):
        """9. SBR dividend = 47.1 cents in the dollar."""
        result = self._calculate()
        assert result["sbr_dividend_cents"] == 47.1

    def test_pbm_sbr_available_for_distribution(self):
        """10. Available = $516,000 - ($516,000 x 0.10) = $464,400."""
        result = self._calculate()
        assert result["sbr_available"] == pytest.approx(464_400.0)

    def test_pbm_liquidation_dividend_zero(self):
        """11. Liquidation dividend = 0 cents (fees exceed recovery)."""
        result = self._calculate()
        assert result["liquidation_dividend_cents"] == 0.0

    def test_pbm_liquidation_asset_recoveries(self):
        """12. Individual asset recoveries at default rates."""
        assets = self._build_pbm_assets()
        by_type = {a["asset_type"]: a for a in assets}

        for asset_type, expected_key in [
            ("cash", "cash"),
            ("receivables", "receivables"),
            ("inventory", "inventory"),
            ("equipment", "equipment"),
            ("loans_related", "loans_related"),
            ("loans_shareholder", "loans_shareholder"),
        ]:
            assert by_type[asset_type]["liquidation_value"] == pytest.approx(
                EXPECTED_LIQUIDATION_RECOVERIES[expected_key], abs=0.01
            ), f"Recovery mismatch for {asset_type}"

    def test_pbm_liquidation_costs_exceed_recoveries(self):
        """13. Total recoveries (~$60k) minus fees ($65k) = negative -> $0 available."""
        assets = self._build_pbm_assets()
        total_recovered = sum(a["liquidation_value"] for a in assets)
        total_fees = (
            PBM_PLAN["est_liquidator_fees"]
            + PBM_PLAN["est_legal_fees"]
            + PBM_PLAN["est_disbursements"]
        )
        assert total_recovered < total_fees
        result = self._calculate()
        assert result["liquidation_available"] == 0.0


# ===========================================================================
#  PAYMENT SCHEDULE REGRESSION (Tests 14-18)
# ===========================================================================


class TestPBMPaymentSchedule:
    """Tests 14-18: Payment schedule generation and verification."""

    def _schedule(self) -> dict:
        return payment_generator.generate(PBM_PLAN)

    def test_pbm_payment_count(self):
        """14. Total payments = 24 (2 initial + 22 ongoing)."""
        schedule = self._schedule()
        assert len(schedule["entries"]) == 24

    def test_pbm_initial_payments(self):
        """15. Payments 1-2: $32,500 each."""
        schedule = self._schedule()
        for i in range(2):
            entry = schedule["entries"][i]
            assert entry["total_payment"] == pytest.approx(32_500.0)
            assert entry["payment_number"] == i + 1

    def test_pbm_ongoing_payments(self):
        """16. Payments 3-24: $20,500 each."""
        schedule = self._schedule()
        for i in range(2, 24):
            entry = schedule["entries"][i]
            assert entry["total_payment"] == pytest.approx(20_500.0)
            assert entry["payment_number"] == i + 1

    def test_pbm_total_contribution(self):
        """17. Sum of all payments = $516,000."""
        schedule = self._schedule()
        total = sum(e["total_payment"] for e in schedule["entries"])
        assert total == pytest.approx(516_000.0)
        assert schedule["total_contribution"] == pytest.approx(516_000.0)

    def test_pbm_fee_allocation(self):
        """18. Each payment splits into net dividend (90%) and RP fee (10%)."""
        schedule = self._schedule()
        # Payment 1: $32,500 -> $29,250 net + $3,250 fee
        p1 = schedule["entries"][0]
        assert p1["net_dividend"] == pytest.approx(29_250.0)
        assert p1["practitioner_fee"] == pytest.approx(3_250.0)

        # Payment 3: $20,500 -> $18,450 net + $2,050 fee
        p3 = schedule["entries"][2]
        assert p3["net_dividend"] == pytest.approx(18_450.0)
        assert p3["practitioner_fee"] == pytest.approx(2_050.0)


# ===========================================================================
#  PII SCRUB REGRESSION (Tests 19-22)
# ===========================================================================


class TestPBMNarrativeScrub:
    """Tests 19-22: PII scrubbing of PBM director narrative."""

    def _scrub_narrative(self):
        narrative = _load_narrative()
        known = {
            "client_name": ["Dr James Mitchell"],
        }
        return scrub(narrative, known_entities=known)

    def test_pbm_narrative_scrub_no_leaks(self):
        """19. Scrub PBM narrative -> ZERO instances of PII in scrubbed output."""
        result = self._scrub_narrative()
        scrubbed = result.scrubbed_text

        # Director name must not appear
        assert "Dr James Mitchell" not in scrubbed
        assert "James Mitchell" not in scrubbed

        # Address must not appear
        assert "42 Harbour Road" not in scrubbed
        assert "Manly NSW 2095" not in scrubbed

        # Trust name must not appear
        assert "Mitchell Family Trust" not in scrubbed

    def test_pbm_narrative_scrub_preserves_amounts(self):
        """20. After scrubbing, dollar amounts must still be present."""
        result = self._scrub_narrative()
        scrubbed = result.scrubbed_text

        # These amounts appear in the narrative text itself
        assert "$573,230.31" in scrubbed
        assert "$516,000" in scrubbed
        assert "$87,500" in scrubbed
        assert "$142,000" in scrubbed
        assert "$412,000" in scrubbed

    def test_pbm_narrative_scrub_preserves_medical_terms(self):
        """21. After scrubbing, medical/industry terms must still be present."""
        result = self._scrub_narrative()
        scrubbed = result.scrubbed_text

        assert "orthopaedic" in scrubbed
        assert "allograft" in scrubbed
        assert "surgical consumables" in scrubbed

    def test_pbm_narrative_scrub_round_trip(self):
        """22. restore(scrub(narrative)) == original narrative."""
        narrative = _load_narrative()
        known = {
            "client_name": ["Dr James Mitchell"],
        }
        result = scrub(narrative, known_entities=known)
        restored = restore(result.scrubbed_text, result.entity_map)
        assert restored == narrative


# ===========================================================================
#  DOCUMENT GENERATION REGRESSION (Tests 23-28)
# ===========================================================================


class TestPBMDocumentGeneration:
    """Tests 23-28: Document generation from PBM data."""

    def _comparison_data(self) -> dict:
        parsed_bs = parser.parse_balance_sheet(PBM_BALANCE_SHEET)
        assets = comparison_engine.build_assets_from_balance_sheet(parsed_bs)
        return comparison_engine.calculate(assets, PBM_CREDITORS_TOTAL, PBM_PLAN)

    def _schedule_data(self) -> dict:
        return payment_generator.generate(PBM_PLAN)

    def _sample_sections(self) -> list[dict]:
        return [
            {"section": "background", "content": "Point Blank Medical Pty Ltd was incorporated in 2018.", "status": "approved"},
            {"section": "expert_advice", "content": "The directors sought professional advice.", "status": "approved"},
            {"section": "plan_summary", "content": "The proposed plan involves $516,000 contribution.", "status": "approved"},
            {"section": "viability", "content": "The company has implemented cost reductions.", "status": "approved"},
            {"section": "comparison_commentary", "content": "Under SBR, creditors receive 47.1 cents.", "status": "approved"},
            {"section": "distress_events", "content": "In mid-2025, the company lost a major contract.", "status": "approved"},
        ]

    def test_pbm_comparison_docx_valid(self):
        """23. Generate comparison .docx from PBM data -> valid Document."""
        data = self._comparison_data()
        filepath = doc_generator.generate_comparison_docx(
            data, PBM_COMPANY_NAME, acn=PBM_ACN
        )
        try:
            assert filepath.exists()
            assert filepath.suffix == ".docx"
            doc = DocxDocument(str(filepath))
            assert len(doc.paragraphs) > 0
        finally:
            if filepath.exists():
                os.remove(str(filepath))

    def test_pbm_comparison_docx_contains_golden_values(self):
        """24. Read back .docx -> verify '47.1' and '464,400' in table cells."""
        data = self._comparison_data()
        filepath = doc_generator.generate_comparison_docx(
            data, PBM_COMPANY_NAME, acn=PBM_ACN
        )
        try:
            doc = DocxDocument(str(filepath))
            text = _all_text(doc)
            assert "47.1" in text
            assert "464,400" in text
        finally:
            if filepath.exists():
                os.remove(str(filepath))

    def test_pbm_payment_schedule_docx_valid(self):
        """25. Generate payment schedule .docx -> valid Document."""
        data = self._schedule_data()
        filepath = doc_generator.generate_payment_schedule_docx(
            data, PBM_COMPANY_NAME
        )
        try:
            assert filepath.exists()
            assert filepath.suffix == ".docx"
            doc = DocxDocument(str(filepath))
            assert len(doc.paragraphs) > 0
        finally:
            if filepath.exists():
                os.remove(str(filepath))

    def test_pbm_payment_schedule_docx_row_count(self):
        """26. Read back -> table has 24 data rows (plus header and total)."""
        data = self._schedule_data()
        filepath = doc_generator.generate_payment_schedule_docx(
            data, PBM_COMPANY_NAME
        )
        try:
            doc = DocxDocument(str(filepath))
            table = doc.tables[0]
            # 1 header + 24 data + 1 total = 26
            assert len(table.rows) == 26
        finally:
            if filepath.exists():
                os.remove(str(filepath))

    def test_pbm_company_statement_docx_valid(self):
        """27. Generate company statement .docx -> valid Document."""
        sections = self._sample_sections()
        filepath = doc_generator.generate_company_statement_docx(
            sections, PBM_COMPANY_NAME, acn=PBM_ACN,
            practitioner_name="Test Practitioner",
        )
        try:
            assert filepath.exists()
            assert filepath.suffix == ".docx"
            doc = DocxDocument(str(filepath))
            assert len(doc.paragraphs) > 0
        finally:
            if filepath.exists():
                os.remove(str(filepath))

    def test_pbm_company_statement_docx_has_all_sections(self):
        """28. Read back -> all 6 section headings present (I through VI)."""
        sections = self._sample_sections()
        filepath = doc_generator.generate_company_statement_docx(
            sections, PBM_COMPANY_NAME, acn=PBM_ACN,
            practitioner_name="Test Practitioner",
        )
        try:
            doc = DocxDocument(str(filepath))
            text = _all_text(doc)
            for numeral in ["I", "II", "III", "IV", "V", "VI"]:
                assert f"SECTION {numeral}" in text, \
                    f"Section {numeral} not found in document"
        finally:
            if filepath.exists():
                os.remove(str(filepath))


# ===========================================================================
#  FULL PIPELINE REGRESSION (Test 29)
# ===========================================================================


class TestPBMFullPipeline:
    """Test 29: The master test — runs the entire flow end-to-end."""

    def test_pbm_full_pipeline(self):
        """
        29. Full pipeline: parse -> creditor schedule -> comparison ->
        payment schedule -> generate all three .docx files.
        """
        generated_files: list[Path] = []

        try:
            # Step 1: Parse aged payables CSV
            creditors_parsed = parser.parse_aged_payables(PBM_AGED_PAYABLES)
            assert len(creditors_parsed) == 6

            # Step 2: Parse balance sheet CSV
            balance_sheet = parser.parse_balance_sheet(PBM_BALANCE_SHEET)
            assert balance_sheet["total_liabilities"] == pytest.approx(PBM_CREDITORS_TOTAL, abs=0.01)

            # Step 3: Build creditor schedule
            creditors = creditor_svc.build_from_parsed(creditors_parsed)
            assert len(creditors) == 6
            totals = creditor_svc.calculate_totals(creditors)
            assert totals["total_claims"] == pytest.approx(
                sum(EXPECTED_CREDITORS.values()), abs=0.01
            )

            # Step 4: Run comparison engine
            assets = comparison_engine.build_assets_from_balance_sheet(balance_sheet)
            assert len(assets) == 6
            comparison = comparison_engine.calculate(
                assets, PBM_CREDITORS_TOTAL, PBM_PLAN
            )
            assert comparison["sbr_dividend_cents"] == 47.1
            assert comparison["liquidation_dividend_cents"] == 0.0
            assert comparison["sbr_available"] == pytest.approx(464_400.0)

            # Step 5: Generate payment schedule
            schedule = payment_generator.generate(PBM_PLAN)
            assert len(schedule["entries"]) == 24
            assert schedule["total_contribution"] == pytest.approx(516_000.0)

            # Step 6: Generate comparison .docx
            comp_path = doc_generator.generate_comparison_docx(
                comparison, PBM_COMPANY_NAME, acn=PBM_ACN
            )
            generated_files.append(comp_path)
            assert comp_path.exists()
            doc = DocxDocument(str(comp_path))
            assert len(doc.tables) >= 1

            # Step 7: Generate payment schedule .docx
            sched_path = doc_generator.generate_payment_schedule_docx(
                schedule, PBM_COMPANY_NAME
            )
            generated_files.append(sched_path)
            assert sched_path.exists()
            doc = DocxDocument(str(sched_path))
            assert len(doc.tables) >= 1

            # Step 8: Generate company statement .docx
            sections = [
                {"section": "background", "content": "PBM background.", "status": "approved"},
                {"section": "expert_advice", "content": "Expert advice.", "status": "approved"},
                {"section": "plan_summary", "content": "Plan summary.", "status": "approved"},
                {"section": "viability", "content": "Viability.", "status": "approved"},
                {"section": "comparison_commentary", "content": "Comparison.", "status": "approved"},
                {"section": "distress_events", "content": "Distress events.", "status": "approved"},
            ]
            stmt_path = doc_generator.generate_company_statement_docx(
                sections, PBM_COMPANY_NAME, acn=PBM_ACN,
                practitioner_name="Test Practitioner",
            )
            generated_files.append(stmt_path)
            assert stmt_path.exists()
            doc = DocxDocument(str(stmt_path))
            assert len(doc.paragraphs) > 0

        finally:
            for f in generated_files:
                if f.exists():
                    os.remove(str(f))
