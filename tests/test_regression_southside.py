"""
Southside Civil Regression Tests
================================
Second test case validating the system works for non-medical companies.
Tests the same pipeline as PBM but with different data and industry.

Golden values (from Southside Civil case):
- Total creditor claims: $742,810.00 (excluding forgiven related party)
- SBR contribution: $380,000
- Practitioner fee: 10%
- SBR available for distribution: $342,000
- SBR dividend: 46.0 cents in the dollar
- Liquidation total recoveries: $201,370
- Liquidation fees: $90,000 ($60,000 liquidator + $30,000 legal)
- Liquidation available: $111,370
- Liquidation dividend: 15.0 cents in the dollar (non-zero!)
- Initial payments: 2 x $25,000
- Ongoing payments: 22 x $15,000
"""

from __future__ import annotations

import os
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from services.file_parser import FileParser
from services.creditor_schedule import CreditorScheduleService
from services.comparison_engine import ComparisonEngine
from services.payment_schedule import PaymentScheduleGenerator
from services.document_generator import DocumentGenerator
from services.privacy_vault import scrub, restore

# ---------------------------------------------------------------------------
#  Fixtures & Constants
# ---------------------------------------------------------------------------

FIXTURES = Path(__file__).parent / "fixtures"

SOUTHSIDE_AGED_PAYABLES = str(FIXTURES / "southside_aged_payables.csv")
SOUTHSIDE_BALANCE_SHEET = str(FIXTURES / "southside_balance_sheet.csv")
SOUTHSIDE_DIRECTOR_NARRATIVE = str(FIXTURES / "southside_director_narrative.txt")

SOUTHSIDE_COMPANY_NAME = "Southside Civil Pty Ltd"
SOUTHSIDE_ACN = "987654321"
SOUTHSIDE_CREDITORS_TOTAL = 742_810.00

SOUTHSIDE_PLAN = {
    "total_contribution": 380_000.0,
    "practitioner_fee_pct": 10.0,
    "num_initial_payments": 2,
    "initial_payment_amount": 25_000.0,
    "num_ongoing_payments": 22,
    "ongoing_payment_amount": 15_000.0,
    "est_liquidator_fees": 60_000.0,
    "est_legal_fees": 30_000.0,
    "est_disbursements": 0.0,
}

# Expected creditor data from aged payables CSV
EXPECTED_CREDITORS = {
    "ATO - ITA": 312_450.00,
    "ATO - ICA": 87_600.00,
    "ATO - SGC": 45_200.00,
    "Coates Hire": 28_750.00,
    "Boral Limited": 67_320.00,
    "Hanson Construction Materials": 41_890.00,
    "WorkCover QLD": 3_200.00,
    "Westpac Equipment Finance": 156_400.00,
    "Torres Family Trust": 85_000.00,
}

# Expected balance sheet asset values
EXPECTED_ASSETS = {
    "cash": 12_450.00,
    "receivables": 145_600.00,
    "inventory": 23_400.00,
    "equipment": 289_000.00,
    "motor_vehicles": 67_500.00,
    "goodwill": 50_000.00,
}

# Practitioner-override recovery rates for construction assets
CUSTOM_RECOVERY_RATES = {
    "equipment": 0.40,  # Construction P&E has higher resale than medical
}

# Expected liquidation recovery at custom rates
EXPECTED_LIQUIDATION_RECOVERIES = {
    "cash": 12_450.00 * 0.20,             # $2,490.00
    "receivables": 145_600.00 * 0.30,     # $43,680.00
    "inventory": 23_400.00 * 0.25,        # $5,850.00
    "equipment": 289_000.00 * 0.40,       # $115,600.00
    "motor_vehicles": 67_500.00 * 0.50,   # $33,750.00
    "goodwill": 50_000.00 * 0.00,         # $0.00
}

# Service singletons
parser = FileParser()
creditor_svc = CreditorScheduleService()
comparison_engine = ComparisonEngine()
payment_generator = PaymentScheduleGenerator()
doc_generator = DocumentGenerator()


# ---------------------------------------------------------------------------
#  Helpers
# ---------------------------------------------------------------------------

def _all_text(doc: DocxDocument) -> str:
    """Extract all paragraph and table text from a docx document."""
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


def _load_narrative() -> str:
    """Load the Southside Civil director narrative fixture."""
    return Path(SOUTHSIDE_DIRECTOR_NARRATIVE).read_text()


def _build_southside_assets() -> list[dict]:
    """Build asset list from balance sheet with practitioner recovery overrides."""
    parsed_bs = parser.parse_balance_sheet(SOUTHSIDE_BALANCE_SHEET)
    assets = comparison_engine.build_assets_from_balance_sheet(parsed_bs)
    # Apply practitioner overrides for construction-specific recovery rates
    for asset in assets:
        if asset["asset_type"] in CUSTOM_RECOVERY_RATES:
            rate = CUSTOM_RECOVERY_RATES[asset["asset_type"]]
            asset["liquidation_recovery_pct"] = rate
            asset["liquidation_value"] = asset["book_value"] * rate
    return assets


def _calculate() -> dict:
    """Run the full comparison engine with Southside data."""
    assets = _build_southside_assets()
    return comparison_engine.calculate(assets, SOUTHSIDE_CREDITORS_TOTAL, SOUTHSIDE_PLAN)


# ===========================================================================
#  1. AGED PAYABLES PARSING
# ===========================================================================


class TestSouthsideAgedPayablesParse:
    """Test 1: Parse aged payables CSV and verify all 9 creditors extracted."""

    def test_southside_aged_payables_parse(self):
        """1. Parse southside_aged_payables.csv -> verify ALL 9 creditors extracted."""
        result = parser.parse_aged_payables(SOUTHSIDE_AGED_PAYABLES)
        assert len(result) == 9
        names = {c["creditor_name"] for c in result}
        assert names == set(EXPECTED_CREDITORS.keys())


# ===========================================================================
#  2. BALANCE SHEET PARSING
# ===========================================================================


class TestSouthsideBalanceSheetParse:
    """Test 2: Parse balance sheet CSV and verify 6 asset categories extracted."""

    def test_southside_balance_sheet_parse(self):
        """2. Parse southside_balance_sheet.csv -> verify 6 asset categories."""
        result = parser.parse_balance_sheet(SOUTHSIDE_BALANCE_SHEET)
        for key, expected_val in EXPECTED_ASSETS.items():
            assert key in result, f"Missing asset category: {key}"
            assert result[key] == pytest.approx(expected_val, abs=0.01), \
                f"{key}: expected {expected_val}, got {result[key]}"


# ===========================================================================
#  3. CREDITOR CLASSIFICATION
# ===========================================================================


class TestSouthsideCreditorClassification:
    """Test 3: Auto-classify Southside creditors into correct categories."""

    def test_southside_creditor_classification(self):
        """3. Auto-classify creditors into correct categories."""
        parsed = parser.parse_aged_payables(SOUTHSIDE_AGED_PAYABLES)
        creditors = creditor_svc.build_from_parsed(parsed)
        by_name = {c["creditor_name"]: c for c in creditors}

        assert by_name["ATO - ITA"]["category"] == "ato_ita"
        assert by_name["ATO - ICA"]["category"] == "ato_ica"
        assert by_name["ATO - SGC"]["category"] == "ato_ica"
        assert by_name["Coates Hire"]["category"] == "trade"
        assert by_name["Boral Limited"]["category"] == "trade"
        assert by_name["Hanson Construction Materials"]["category"] == "trade"
        assert by_name["WorkCover QLD"]["category"] == "workers_comp"
        assert by_name["Westpac Equipment Finance"]["category"] == "finance"
        assert by_name["Torres Family Trust"]["category"] == "trade"


# ===========================================================================
#  4. RELATED PARTY FLAGGING
# ===========================================================================


class TestSouthsideRelatedParty:
    """Test 4: Torres Family Trust flagged as related party."""

    def test_southside_related_party(self):
        """4. Torres Family Trust flagged -> can_vote=False."""
        parsed = parser.parse_aged_payables(SOUTHSIDE_AGED_PAYABLES)
        creditors = creditor_svc.build_from_parsed(parsed)
        by_name = {c["creditor_name"]: c for c in creditors}

        torres = by_name["Torres Family Trust"]
        creditor_svc.flag_related_party(torres, True)
        assert torres["is_related_party"] is True
        assert torres["can_vote"] is False


# ===========================================================================
#  5. DEBT FORGIVENESS
# ===========================================================================


class TestSouthsideDebtForgiveness:
    """Test 5: Torres Family Trust forgiven -> excluded from admitted totals."""

    def test_southside_debt_forgiveness(self):
        """5. Torres Family Trust marked as forgiven -> excluded from admitted."""
        parsed = parser.parse_aged_payables(SOUTHSIDE_AGED_PAYABLES)
        creditors = creditor_svc.build_from_parsed(parsed)
        by_name = {c["creditor_name"]: c for c in creditors}

        torres = by_name["Torres Family Trust"]
        creditor_svc.flag_related_party(torres, True)
        creditor_svc.update_status(torres, "forgiven")

        totals = creditor_svc.calculate_totals(creditors)
        assert totals["total_excluded"] == 1
        # Torres Family Trust's $85,000 should NOT be in total_admitted
        assert totals["total_admitted"] == pytest.approx(
            sum(EXPECTED_CREDITORS.values()) - EXPECTED_CREDITORS["Torres Family Trust"],
            abs=0.01,
        )


# ===========================================================================
#  6. TOTAL CLAIMS
# ===========================================================================


class TestSouthsideTotalClaims:
    """Test 6: Total claims = $742,810."""

    def test_southside_total_claims(self):
        """6. Total claims from balance sheet = $742,810.00."""
        result = parser.parse_balance_sheet(SOUTHSIDE_BALANCE_SHEET)
        assert result["total_liabilities"] == pytest.approx(SOUTHSIDE_CREDITORS_TOTAL, abs=0.01)


# ===========================================================================
#  7. SBR DIVIDEND
# ===========================================================================


class TestSouthsideSBRDividend:
    """Test 7: SBR dividend = 46.0 cents in the dollar."""

    def test_southside_sbr_dividend(self):
        """7. SBR dividend = 46.0 cents in the dollar."""
        result = _calculate()
        assert result["sbr_dividend_cents"] == 46.0
        assert result["sbr_available"] == pytest.approx(342_000.0)


# ===========================================================================
#  8. LIQUIDATION DIVIDEND
# ===========================================================================


class TestSouthsideLiquidationDividend:
    """Test 8: Liquidation dividend = 15.0 cents (non-zero, unlike PBM)."""

    def test_southside_liquidation_dividend(self):
        """8. Liquidation dividend = 15.0 cents in the dollar."""
        result = _calculate()
        assert result["liquidation_dividend_cents"] == 15.0


# ===========================================================================
#  9. LIQUIDATION POSITIVE
# ===========================================================================


class TestSouthsideLiquidationPositive:
    """Test 9: Liquidation available > $0 (unlike PBM where it was $0)."""

    def test_southside_liquidation_positive(self):
        """9. Liquidation available = $111,370 > $0."""
        result = _calculate()
        assert result["liquidation_available"] > 0
        assert result["liquidation_available"] == pytest.approx(111_370.0, abs=1.0)


# ===========================================================================
#  10. PAYMENT COUNT
# ===========================================================================


class TestSouthsidePaymentCount:
    """Test 10: 24 payments (2 initial + 22 ongoing)."""

    def test_southside_payment_count(self):
        """10. Total payments = 24."""
        schedule = payment_generator.generate(SOUTHSIDE_PLAN)
        assert len(schedule["entries"]) == 24


# ===========================================================================
#  11. PAYMENT TOTAL
# ===========================================================================


class TestSouthsidePaymentTotal:
    """Test 11: Sum of all payments = $380,000."""

    def test_southside_payment_total(self):
        """11. Sum of all payments = $380,000."""
        schedule = payment_generator.generate(SOUTHSIDE_PLAN)
        total = sum(e["total_payment"] for e in schedule["entries"])
        assert total == pytest.approx(380_000.0)
        assert schedule["total_contribution"] == pytest.approx(380_000.0)


# ===========================================================================
#  12. COMPARISON DOCX
# ===========================================================================


class TestSouthsideComparisonDocx:
    """Test 12: Generate valid comparison .docx."""

    def test_southside_comparison_docx(self):
        """12. Generate comparison .docx -> valid Document with golden values."""
        data = _calculate()
        filepath = doc_generator.generate_comparison_docx(
            data, SOUTHSIDE_COMPANY_NAME, acn=SOUTHSIDE_ACN
        )
        try:
            assert filepath.exists()
            assert filepath.suffix == ".docx"
            doc = DocxDocument(str(filepath))
            text = _all_text(doc)
            assert "46.0" in text
            assert "342,000" in text
            assert "15.0" in text
        finally:
            if filepath.exists():
                os.remove(str(filepath))


# ===========================================================================
#  13. PAYMENT SCHEDULE DOCX
# ===========================================================================


class TestSouthsidePaymentScheduleDocx:
    """Test 13: Generate valid payment schedule .docx."""

    def test_southside_payment_schedule_docx(self):
        """13. Generate payment schedule .docx -> valid Document."""
        schedule = payment_generator.generate(SOUTHSIDE_PLAN)
        filepath = doc_generator.generate_payment_schedule_docx(
            schedule, SOUTHSIDE_COMPANY_NAME
        )
        try:
            assert filepath.exists()
            assert filepath.suffix == ".docx"
            doc = DocxDocument(str(filepath))
            table = doc.tables[0]
            # 1 header + 24 data + 1 total = 26 rows
            assert len(table.rows) == 26
        finally:
            if filepath.exists():
                os.remove(str(filepath))


# ===========================================================================
#  14. PII SCRUB
# ===========================================================================


class TestSouthsidePIIScrub:
    """Test 14: Director name, address, and trust scrubbed from narrative."""

    def test_southside_pii_scrub(self):
        """14. Scrub Southside narrative -> ZERO instances of PII."""
        narrative = _load_narrative()
        known = {
            "client_name": ["Michael Torres", "Mr Torres"],
        }
        result = scrub(narrative, known_entities=known)
        scrubbed = result.scrubbed_text

        # Director name must not appear
        assert "Michael Torres" not in scrubbed
        assert "Mr Torres" not in scrubbed

        # Address must not appear
        assert "15 Industrial Drive" not in scrubbed
        assert "Rocklea QLD 4106" not in scrubbed

        # Trust name must not appear
        assert "Torres Family Trust" not in scrubbed


# ===========================================================================
#  15. PII PRESERVES AMOUNTS
# ===========================================================================


class TestSouthsidePIIPreservesAmounts:
    """Test 15: Dollar amounts preserved after scrubbing."""

    def test_southside_pii_preserves_amounts(self):
        """15. After scrubbing, dollar amounts must still be present."""
        narrative = _load_narrative()
        known = {
            "client_name": ["Michael Torres", "Mr Torres"],
        }
        result = scrub(narrative, known_entities=known)
        scrubbed = result.scrubbed_text

        assert "$312,450" in scrubbed
        assert "$87,600" in scrubbed
        assert "$45,200" in scrubbed
        assert "$185,000" in scrubbed
        assert "$85,000" in scrubbed
        assert "$380,000" in scrubbed
        assert "$45,000" in scrubbed


# ===========================================================================
#  16. FULL PIPELINE
# ===========================================================================


class TestSouthsideFullPipeline:
    """Test 16: End-to-end parse -> compare -> generate .docx."""

    def test_southside_full_pipeline(self):
        """
        16. Full pipeline: parse -> creditor schedule -> comparison ->
        payment schedule -> generate .docx files.
        """
        generated_files: list[Path] = []

        try:
            # Step 1: Parse aged payables CSV
            creditors_parsed = parser.parse_aged_payables(SOUTHSIDE_AGED_PAYABLES)
            assert len(creditors_parsed) == 9

            # Step 2: Parse balance sheet CSV
            balance_sheet = parser.parse_balance_sheet(SOUTHSIDE_BALANCE_SHEET)
            assert balance_sheet["total_liabilities"] == pytest.approx(
                SOUTHSIDE_CREDITORS_TOTAL, abs=0.01
            )

            # Step 3: Build creditor schedule and apply forgiveness
            creditors = creditor_svc.build_from_parsed(creditors_parsed)
            assert len(creditors) == 9
            by_name = {c["creditor_name"]: c for c in creditors}

            # Flag Torres Family Trust as related party and forgiven
            torres = by_name["Torres Family Trust"]
            creditor_svc.flag_related_party(torres, True)
            creditor_svc.update_status(torres, "forgiven")

            totals = creditor_svc.calculate_totals(creditors)
            assert totals["total_admitted"] == pytest.approx(
                SOUTHSIDE_CREDITORS_TOTAL, abs=0.01
            )

            # Step 4: Run comparison engine with custom recovery rates
            assets = _build_southside_assets()
            # Verify we got the expected asset types (excluding goodwill at $0 recovery)
            asset_types = {a["asset_type"] for a in assets}
            assert "cash" in asset_types
            assert "receivables" in asset_types
            assert "equipment" in asset_types
            assert "motor_vehicles" in asset_types

            comparison = comparison_engine.calculate(
                assets, SOUTHSIDE_CREDITORS_TOTAL, SOUTHSIDE_PLAN
            )
            assert comparison["sbr_dividend_cents"] == 46.0
            assert comparison["liquidation_dividend_cents"] == 15.0
            assert comparison["sbr_available"] == pytest.approx(342_000.0)
            assert comparison["liquidation_available"] == pytest.approx(111_370.0, abs=1.0)

            # Step 5: Generate payment schedule
            schedule = payment_generator.generate(SOUTHSIDE_PLAN)
            assert len(schedule["entries"]) == 24
            assert schedule["total_contribution"] == pytest.approx(380_000.0)

            # Step 6: Generate comparison .docx
            comp_path = doc_generator.generate_comparison_docx(
                comparison, SOUTHSIDE_COMPANY_NAME, acn=SOUTHSIDE_ACN
            )
            generated_files.append(comp_path)
            assert comp_path.exists()
            doc = DocxDocument(str(comp_path))
            assert len(doc.tables) >= 1

            # Step 7: Generate payment schedule .docx
            sched_path = doc_generator.generate_payment_schedule_docx(
                schedule, SOUTHSIDE_COMPANY_NAME
            )
            generated_files.append(sched_path)
            assert sched_path.exists()
            doc = DocxDocument(str(sched_path))
            assert len(doc.tables) >= 1

        finally:
            for f in generated_files:
                if f.exists():
                    os.remove(str(f))
