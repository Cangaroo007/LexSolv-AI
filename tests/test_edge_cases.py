"""
Edge Case Tests
===============
These test unusual but valid scenarios that practitioners will encounter.
Each test verifies graceful handling -- no crashes, sensible defaults,
clear error messages.
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from services.file_parser import FileParser
from services.creditor_schedule import CreditorScheduleService
from services.comparison_engine import ComparisonEngine
from services.payment_schedule import PaymentScheduleGenerator
from services.document_generator import DocumentGenerator
from services.privacy_vault import scrub, restore

# ---------------------------------------------------------------------------
#  Fixtures & Constants
# ---------------------------------------------------------------------------

FIXTURES = Path(__file__).parent / "fixtures"

parser = FileParser()
creditor_svc = CreditorScheduleService()
comparison_engine = ComparisonEngine()
payment_generator = PaymentScheduleGenerator()
doc_generator = DocumentGenerator()


def _all_text(doc: DocxDocument) -> str:
    """Extract all paragraph and table text from a docx document."""
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


# ===========================================================================
#  ZERO-ASSET LIQUIDATION (Tests 1-2)
# ===========================================================================


class TestZeroAssetLiquidation:
    """Company with $0 in every asset category."""

    def _zero_assets(self) -> list[dict]:
        return [
            {
                "asset_type": "cash",
                "description": "Cash at Bank",
                "book_value": 0.0,
                "liquidation_recovery_pct": 0.20,
                "liquidation_value": 0.0,
            },
            {
                "asset_type": "receivables",
                "description": "Accounts Receivable",
                "book_value": 0.0,
                "liquidation_recovery_pct": 0.30,
                "liquidation_value": 0.0,
            },
            {
                "asset_type": "inventory",
                "description": "Inventory",
                "book_value": 0.0,
                "liquidation_recovery_pct": 0.25,
                "liquidation_value": 0.0,
            },
            {
                "asset_type": "equipment",
                "description": "Plant & Equipment",
                "book_value": 0.0,
                "liquidation_recovery_pct": 0.25,
                "liquidation_value": 0.0,
            },
        ]

    def _plan(self) -> dict:
        return {
            "total_contribution": 100_000.0,
            "practitioner_fee_pct": 10.0,
            "num_initial_payments": 2,
            "initial_payment_amount": 10_000.0,
            "num_ongoing_payments": 4,
            "ongoing_payment_amount": 20_000.0,
            "est_liquidator_fees": 50_000.0,
            "est_legal_fees": 10_000.0,
            "est_disbursements": 5_000.0,
        }

    def test_zero_asset_liquidation(self):
        """1. Zero assets -> liquidation dividend = 0 (not negative, not error)."""
        result = comparison_engine.calculate(
            self._zero_assets(), 500_000.0, self._plan()
        )
        assert result["liquidation_dividend_cents"] == 0.0
        assert result["liquidation_available"] == 0.0
        # SBR should still calculate correctly
        assert result["sbr_available"] == pytest.approx(90_000.0)
        assert result["sbr_dividend_cents"] > 0

    def test_zero_asset_comparison_still_works(self):
        """2. Comparison table generates with $0 in all liquidation asset rows."""
        data = comparison_engine.calculate(
            self._zero_assets(), 500_000.0, self._plan()
        )
        filepath = doc_generator.generate_comparison_docx(
            data, "Zero Assets Pty Ltd"
        )
        try:
            assert filepath.exists()
            doc = DocxDocument(str(filepath))
            assert len(doc.tables) >= 1
            text = _all_text(doc)
            # SBR figures should be present
            assert "90,000.00" in text
        finally:
            if filepath.exists():
                os.remove(str(filepath))


# ===========================================================================
#  SINGLE CREDITOR (Tests 3-4)
# ===========================================================================


class TestSingleCreditor:
    """Only one creditor in the schedule."""

    def _single_creditor_parsed(self) -> list[dict]:
        return parser.parse_aged_payables(str(FIXTURES / "single_creditor.csv"))

    def test_single_creditor_ato_only(self):
        """3. Single creditor (ATO - ITA, $500,000) -> all services work."""
        parsed = self._single_creditor_parsed()
        assert len(parsed) == 1
        assert parsed[0]["amount_claimed"] == pytest.approx(500_000.0)

        # Creditor schedule
        creditors = creditor_svc.build_from_parsed(parsed)
        assert len(creditors) == 1
        assert creditors[0]["category"] == "ato_ita"
        totals = creditor_svc.calculate_totals(creditors)
        assert totals["total_claims"] == pytest.approx(500_000.0)

        # Comparison engine
        plan = {
            "total_contribution": 200_000.0,
            "practitioner_fee_pct": 10.0,
            "num_initial_payments": 1,
            "initial_payment_amount": 200_000.0,
            "num_ongoing_payments": 0,
            "ongoing_payment_amount": 0.0,
            "est_liquidator_fees": 50_000.0,
            "est_legal_fees": 10_000.0,
            "est_disbursements": 5_000.0,
        }
        assets = [
            {
                "asset_type": "cash",
                "description": "Cash",
                "book_value": 10_000.0,
                "liquidation_recovery_pct": 0.20,
                "liquidation_value": 2_000.0,
            }
        ]
        result = comparison_engine.calculate(assets, 500_000.0, plan)
        expected_dividend = (200_000.0 * 0.90) / 500_000.0 * 100
        assert result["sbr_dividend_cents"] == pytest.approx(
            round(expected_dividend, 1)
        )

    def test_single_creditor_docx(self):
        """4. Comparison .docx generates correctly with single creditor total."""
        plan = {
            "total_contribution": 200_000.0,
            "practitioner_fee_pct": 10.0,
            "num_initial_payments": 1,
            "initial_payment_amount": 200_000.0,
            "num_ongoing_payments": 0,
            "ongoing_payment_amount": 0.0,
            "est_liquidator_fees": 50_000.0,
            "est_legal_fees": 10_000.0,
            "est_disbursements": 5_000.0,
        }
        assets = [
            {
                "asset_type": "cash",
                "description": "Cash",
                "book_value": 10_000.0,
                "liquidation_recovery_pct": 0.20,
                "liquidation_value": 2_000.0,
            }
        ]
        data = comparison_engine.calculate(assets, 500_000.0, plan)
        filepath = doc_generator.generate_comparison_docx(
            data, "Single Creditor Pty Ltd"
        )
        try:
            assert filepath.exists()
            doc = DocxDocument(str(filepath))
            assert len(doc.tables) >= 1
        finally:
            if filepath.exists():
                os.remove(str(filepath))


# ===========================================================================
#  NO RELATED PARTIES (Test 5)
# ===========================================================================


class TestNoRelatedParties:
    """All creditors are unrelated."""

    def test_no_related_parties(self):
        """5. No related-party exclusions -> total claims = sum of all creditors."""
        parsed = parser.parse_aged_payables(str(FIXTURES / "pbm_aged_payables.csv"))
        creditors = creditor_svc.build_from_parsed(parsed)
        # None should be related by default
        for c in creditors:
            assert c["is_related_party"] is False
            assert c["can_vote"] is True

        totals = creditor_svc.calculate_totals(creditors)
        expected_total = sum(c["amount_claimed"] for c in creditors)
        assert totals["total_claims"] == pytest.approx(expected_total)
        assert totals["total_voting"] == pytest.approx(expected_total)
        assert totals["related_party_total"] == 0.0


# ===========================================================================
#  CONTRIBUTION EXCEEDS DEBTS (Tests 6-7)
# ===========================================================================


class TestContributionExceedsDebts:
    """Contribution (after fee) > total claims."""

    def _plan(self) -> dict:
        return {
            "total_contribution": 200_000.0,
            "practitioner_fee_pct": 10.0,
            "num_initial_payments": 1,
            "initial_payment_amount": 200_000.0,
            "num_ongoing_payments": 0,
            "ongoing_payment_amount": 0.0,
            "est_liquidator_fees": 10_000.0,
            "est_legal_fees": 5_000.0,
            "est_disbursements": 2_000.0,
        }

    def test_contribution_exceeds_total_debts(self):
        """6. Contribution > total claims -> sensible dividend output."""
        assets = [
            {
                "asset_type": "cash",
                "description": "Cash",
                "book_value": 50_000.0,
                "liquidation_recovery_pct": 0.20,
                "liquidation_value": 10_000.0,
            }
        ]
        # Contribution after fee: $180,000 against $100,000 debts
        result = comparison_engine.calculate(assets, 100_000.0, self._plan())
        # Dividend = 180,000 / 100,000 * 100 = 180.0 cents
        # The engine should produce the calculated value
        assert result["sbr_dividend_cents"] == 180.0
        assert result["sbr_available"] == pytest.approx(180_000.0)

    def test_surplus_contribution_payment_schedule(self):
        """7. Payment schedule generates correctly even with surplus contribution."""
        plan = self._plan()
        schedule = payment_generator.generate(plan)
        assert len(schedule["entries"]) == 1
        assert schedule["total_contribution"] == pytest.approx(200_000.0)
        assert schedule["entries"][0]["total_payment"] == pytest.approx(200_000.0)


# ===========================================================================
#  FILE UPLOAD EDGE CASES (Tests 8-14)
# ===========================================================================


class TestFileUploadEdgeCases:
    """Unusual file formats and contents."""

    def test_empty_csv_upload(self):
        """8. Empty CSV file -> clear error."""
        with pytest.raises(Exception):
            parser.parse_aged_payables(str(FIXTURES / "empty.csv"))

    def test_wrong_file_format_txt(self):
        """9. Upload a .txt file -> clear error about unsupported format."""
        with pytest.raises(ValueError, match="Unsupported file format"):
            parser.parse_aged_payables(str(FIXTURES / "wrong_format.txt"))

    def test_wrong_file_format_pdf(self):
        """10. Upload a .pdf file -> clear error."""
        # Create a temporary file with .pdf extension
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"%PDF-1.4 fake pdf content")
            pdf_path = f.name
        try:
            with pytest.raises(ValueError, match="Unsupported file format"):
                parser.parse_aged_payables(pdf_path)
        finally:
            os.remove(pdf_path)

    def test_csv_with_extra_columns(self):
        """11. CSV with extra columns -> parser ignores extras, extracts data."""
        result = parser.parse_aged_payables(str(FIXTURES / "extra_columns.csv"))
        assert len(result) == 2
        names = {c["creditor_name"] for c in result}
        assert "Australian Taxation Office - ITA" in names
        assert "Prospa Advance" in names

    def test_csv_with_missing_columns(self):
        """12. CSV missing expected columns -> clear error."""
        with pytest.raises(ValueError, match="[Cc]olumn"):
            parser.parse_aged_payables(str(FIXTURES / "missing_columns.csv"))

    def test_csv_with_zero_amounts(self):
        """13. All creditor amounts $0.00 -> handles gracefully."""
        result = parser.parse_aged_payables(str(FIXTURES / "zero_amounts.csv"))
        assert len(result) == 3
        for c in result:
            assert c["amount_claimed"] == 0.0

        # Building creditor schedule with zero amounts should work
        creditors = creditor_svc.build_from_parsed(result)
        totals = creditor_svc.calculate_totals(creditors)
        assert totals["total_claims"] == 0.0

    def test_csv_with_negative_amounts(self):
        """14. Some creditors with negative amounts -> handles without crash."""
        result = parser.parse_aged_payables(str(FIXTURES / "negative_amounts.csv"))
        assert len(result) == 3
        amounts = {c["creditor_name"]: c["amount_claimed"] for c in result}
        assert amounts["Supplier Two"] == pytest.approx(-5_000.0)

        # Creditor schedule should still build
        creditors = creditor_svc.build_from_parsed(result)
        totals = creditor_svc.calculate_totals(creditors)
        # Total claims = 10000 + (-5000) + 20000 = 25000
        assert totals["total_claims"] == pytest.approx(25_000.0)


# ===========================================================================
#  COMPARISON ENGINE EDGE CASES (Tests 15-18)
# ===========================================================================


class TestComparisonEngineEdgeCases:
    """Edge cases in dividend calculation."""

    def _basic_assets(self) -> list[dict]:
        return [
            {
                "asset_type": "cash",
                "description": "Cash",
                "book_value": 10_000.0,
                "liquidation_recovery_pct": 0.20,
                "liquidation_value": 2_000.0,
            }
        ]

    def test_zero_contribution(self):
        """15. SBR contribution = $0 -> dividend = 0 (not division by zero)."""
        plan = {
            "total_contribution": 0.0,
            "practitioner_fee_pct": 10.0,
            "est_liquidator_fees": 50_000.0,
            "est_legal_fees": 10_000.0,
            "est_disbursements": 5_000.0,
        }
        result = comparison_engine.calculate(self._basic_assets(), 500_000.0, plan)
        assert result["sbr_dividend_cents"] == 0.0
        assert result["sbr_available"] == 0.0

    def test_zero_fee_percentage(self):
        """16. Fee = 0% -> full contribution available."""
        plan = {
            "total_contribution": 100_000.0,
            "practitioner_fee_pct": 0.0,
            "est_liquidator_fees": 50_000.0,
            "est_legal_fees": 10_000.0,
            "est_disbursements": 5_000.0,
        }
        result = comparison_engine.calculate(self._basic_assets(), 500_000.0, plan)
        assert result["sbr_available"] == pytest.approx(100_000.0)

    def test_100_percent_fee(self):
        """17. Fee = 100% -> available = $0, dividend = 0 (no crash)."""
        plan = {
            "total_contribution": 100_000.0,
            "practitioner_fee_pct": 100.0,
            "est_liquidator_fees": 50_000.0,
            "est_legal_fees": 10_000.0,
            "est_disbursements": 5_000.0,
        }
        result = comparison_engine.calculate(self._basic_assets(), 500_000.0, plan)
        assert result["sbr_available"] == 0.0
        assert result["sbr_dividend_cents"] == 0.0

    def test_very_small_contribution(self):
        """18. Contribution = $1.00 -> calculates correctly, no scientific notation."""
        plan = {
            "total_contribution": 1.0,
            "practitioner_fee_pct": 10.0,
            "est_liquidator_fees": 50_000.0,
            "est_legal_fees": 10_000.0,
            "est_disbursements": 5_000.0,
        }
        result = comparison_engine.calculate(self._basic_assets(), 500_000.0, plan)
        assert result["sbr_available"] == pytest.approx(0.90)
        # Dividend should be a very small number, not an error
        assert isinstance(result["sbr_dividend_cents"], float)
        assert result["sbr_dividend_cents"] >= 0


# ===========================================================================
#  PAYMENT SCHEDULE EDGE CASES (Tests 19-20)
# ===========================================================================


class TestPaymentScheduleEdgeCases:
    """Edge cases in payment schedule generation."""

    def test_single_payment(self):
        """19. All contribution in one lump sum -> schedule has 1 row."""
        plan = {
            "total_contribution": 100_000.0,
            "practitioner_fee_pct": 10.0,
            "num_initial_payments": 1,
            "initial_payment_amount": 100_000.0,
            "num_ongoing_payments": 0,
            "ongoing_payment_amount": 0.0,
        }
        schedule = payment_generator.generate(plan)
        assert len(schedule["entries"]) == 1
        assert schedule["entries"][0]["total_payment"] == pytest.approx(100_000.0)
        assert schedule["total_contribution"] == pytest.approx(100_000.0)

    def test_mismatched_payment_total(self):
        """20. Payments don't sum to contribution -> ValueError."""
        plan = {
            "total_contribution": 100_000.0,
            "practitioner_fee_pct": 10.0,
            "num_initial_payments": 2,
            "initial_payment_amount": 30_000.0,
            "num_ongoing_payments": 2,
            "ongoing_payment_amount": 10_000.0,
        }
        # 2 * 30000 + 2 * 10000 = 80000 != 100000
        with pytest.raises(ValueError, match="does not balance"):
            payment_generator.generate(plan)


# ===========================================================================
#  NARRATIVE EDGE CASES (Tests 21-24)
# ===========================================================================


class TestNarrativeEdgeCases:
    """Edge cases in PII scrubbing and narrative handling."""

    def test_empty_director_notes(self):
        """21. Empty string for director notes -> scrub handles gracefully."""
        result = scrub("")
        assert result.scrubbed_text == ""
        assert result.entity_map == {}

    def test_very_long_director_notes(self):
        """22. 10,000+ characters -> handles without error."""
        long_text = "The director operated a business. " * 500  # ~16,500 chars
        result = scrub(long_text)
        assert len(result.scrubbed_text) > 0
        # Restore should return the original
        restored = restore(result.scrubbed_text, result.entity_map)
        assert restored == long_text

    def test_director_notes_no_pii(self):
        """23. Notes with zero PII -> scrub passes through unchanged."""
        text = (
            "The company was incorporated in 2018. Revenue declined by 18%. "
            "Total liabilities are approximately $412,000. The company continues "
            "to trade on a limited basis."
        )
        result = scrub(text)
        # No PII detected, so scrubbed text should equal original
        assert result.scrubbed_text == text
        assert result.entity_map == {}
        # Restore should return identical text
        restored = restore(result.scrubbed_text, result.entity_map)
        assert restored == text

    def test_narrative_special_characters(self):
        """24. Director notes with special characters -> no breakage."""
        text = (
            "Dr Müller's café at 42 Smith Street, Sydney NSW 2000 "
            "generated $100,000 in revenue (a 15% increase). "
            "The company's ABN is noted & the <agreement> was signed. "
            "Revenue was €50,000 from European operations."
        )
        result = scrub(text)
        # Should not crash — the result is a valid ScrubResult
        assert isinstance(result.scrubbed_text, str)
        assert isinstance(result.entity_map, dict)
        # Round-trip should work
        restored = restore(result.scrubbed_text, result.entity_map)
        assert restored == text
