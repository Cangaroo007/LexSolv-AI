#!/usr/bin/env python3
"""
Generate synthetic test fixtures for the Document Ingester test suite.

Run once:  python tests/fixtures/generate_fixtures.py

Produces:
  tests/fixtures/pbm_bank_statement.pdf   — synthetic bank-statement PDF
  tests/fixtures/pbm_balance_sheet.docx   — Word doc of PBM balance sheet
  tests/fixtures/scanned_placeholder.jpg  — low-res JPEG simulating a scan
  tests/fixtures/sparse_scanned.pdf       — PDF with only an image (no text)
  tests/fixtures/nonstandard_excel.xlsx   — Excel with unusual columns
  tests/fixtures/nonstandard.csv          — CSV with unknown columns
"""

from __future__ import annotations

import io
from pathlib import Path

FIXTURE_DIR = Path(__file__).parent


def generate_bank_statement_pdf() -> None:
    """Create a synthetic bank-statement PDF with tables using reportlab."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    out = FIXTURE_DIR / "pbm_bank_statement.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("PBM Pty Ltd — Bank Statement", styles["Title"]))
    elements.append(Spacer(1, 10 * mm))

    # Account details
    account_data = [
        ["Account Name", "PBM Pty Ltd"],
        ["BSB", "062-000"],
        ["Account Number", "1234 5678"],
        ["Statement Period", "01 Jan 2024 – 31 Jan 2024"],
    ]
    t = Table(account_data, colWidths=[50 * mm, 80 * mm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 8 * mm))

    # Transactions
    tx_data = [
        ["Date", "Description", "Debit", "Credit", "Balance"],
        ["01/01/2024", "Opening Balance", "", "", "$62,100.00"],
        ["05/01/2024", "BPAY — ATO ITA", "$3,500.00", "", "$58,600.00"],
        ["12/01/2024", "Deposit — Client Payment", "", "$2,500.00", "$61,100.00"],
        ["20/01/2024", "Transfer — Supplier", "$1,410.73", "", "$59,689.27"],
        ["31/01/2024", "Closing Balance", "", "", "$59,689.27"],
    ]
    t2 = Table(tx_data, colWidths=[28 * mm, 55 * mm, 28 * mm, 28 * mm, 30 * mm])
    t2.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ALIGN", (2, 1), (-1, -1), "RIGHT"),
    ]))
    elements.append(t2)

    doc.build(elements)
    print(f"  ✓ {out.name}")


def generate_balance_sheet_docx() -> None:
    """Create a Word doc version of PBM balance sheet."""
    import docx

    out = FIXTURE_DIR / "pbm_balance_sheet.docx"
    doc = docx.Document()
    doc.core_properties.title = "PBM Balance Sheet"
    doc.core_properties.author = "LexSolv AI"

    doc.add_heading("PBM Pty Ltd — Balance Sheet", level=1)
    doc.add_heading("As at 31 January 2024", level=2)

    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Account"
    headers[1].text = "Amount"

    rows_data = [
        ("Cash at Bank", "$59,689.27"),
        ("Accounts Receivable", "$69,553.24"),
        ("Inventory", "$51,826.62"),
        ("Loans to Related Entities", "$34,964.83"),
        ("Shareholder Loans", "$2,010,000.00"),
        ("Plant & Equipment", "$15,000.00"),
        ("Total Current Liabilities", "$985,777.37"),
    ]
    for i, (acct, amt) in enumerate(rows_data, start=1):
        table.rows[i].cells[0].text = acct
        table.rows[i].cells[1].text = amt

    doc.save(str(out))
    print(f"  ✓ {out.name}")


def generate_scanned_placeholder() -> None:
    """Create a low-res JPEG that simulates a scanned table."""
    from PIL import Image, ImageDraw, ImageFont

    out = FIXTURE_DIR / "scanned_placeholder.jpg"
    img = Image.new("RGB", (800, 400), "white")
    draw = ImageDraw.Draw(img)

    # Draw a simple table grid
    for y in range(50, 351, 50):
        draw.line([(50, y), (750, y)], fill="black", width=1)
    for x in [50, 250, 500, 750]:
        draw.line([(x, 50), (x, 350)], fill="black", width=1)

    # Add text (use default font)
    try:
        font = ImageFont.load_default()
    except Exception:
        font = None

    cells = [
        (60, 55, "Account"), (260, 55, "Debit"), (510, 55, "Credit"),
        (60, 105, "Cash at Bank"), (260, 105, "59,689.27"), (510, 105, ""),
        (60, 155, "Receivables"), (260, 155, "69,553.24"), (510, 155, ""),
        (60, 205, "Inventory"), (260, 205, "51,826.62"), (510, 205, ""),
        (60, 255, "Liabilities"), (260, 255, ""), (510, 255, "985,777.37"),
    ]
    for x, y, text in cells:
        draw.text((x, y), text, fill="black", font=font)

    img.save(str(out), "JPEG", quality=60)
    print(f"  ✓ {out.name}")


def generate_sparse_scanned_pdf() -> None:
    """Create a PDF that contains only an image — no extractable text."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from PIL import Image

    out = FIXTURE_DIR / "sparse_scanned.pdf"

    # Create a tiny image to embed
    img = Image.new("RGB", (200, 100), "lightgrey")
    img_buf = io.BytesIO()
    img.save(img_buf, format="JPEG")
    img_path = FIXTURE_DIR / "_temp_scan_img.jpg"
    img_path.write_bytes(img_buf.getvalue())

    c = canvas.Canvas(str(out), pagesize=A4)
    c.drawImage(str(img_path), 100, 500, width=300, height=150)
    c.showPage()
    c.save()

    img_path.unlink()  # clean up temp image
    print(f"  ✓ {out.name}")


def generate_nonstandard_excel() -> None:
    """Create an Excel file with non-standard columns across multiple sheets."""
    import pandas as pd

    out = FIXTURE_DIR / "nonstandard_excel.xlsx"
    with pd.ExcelWriter(str(out), engine="openpyxl") as writer:
        df1 = pd.DataFrame({
            "Item Code": ["A001", "A002", "A003"],
            "Widget Name": ["Flange", "Gasket", "Bolt"],
            "Qty On Hand": [150, 42, 800],
            "Unit Cost": [12.50, 3.75, 0.45],
        })
        df1.to_excel(writer, sheet_name="Inventory", index=False)

        df2 = pd.DataFrame({
            "Employee": ["Smith J", "Doe A"],
            "Hours Worked": [38, 42],
            "Rate": [45.00, 52.00],
        })
        df2.to_excel(writer, sheet_name="Payroll", index=False)

    print(f"  ✓ {out.name}")


def generate_nonstandard_csv() -> None:
    """Create a CSV with completely unknown column names."""
    out = FIXTURE_DIR / "nonstandard.csv"
    out.write_text(
        "SKU,Product,Warehouse,Bin,Stock Level\n"
        "W-001,Widget A,Sydney,B3,120\n"
        "W-002,Widget B,Melbourne,A1,85\n"
        "W-003,Widget C,Brisbane,C7,200\n"
    )
    print(f"  ✓ {out.name}")


if __name__ == "__main__":
    print("Generating test fixtures …")
    generate_bank_statement_pdf()
    generate_balance_sheet_docx()
    generate_scanned_placeholder()
    generate_sparse_scanned_pdf()
    generate_nonstandard_excel()
    generate_nonstandard_csv()
    print("Done.")
