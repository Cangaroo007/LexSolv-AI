"""
LexSolv AI — Statutory Document Generator.

Generates draft insolvency documents in .docx format:
  - DIRRI (Declaration of Independence, Relevant Relationships and Indemnities)
  - Safe Harbour checklist (future)

Uses the ``python-docx`` library to build professional Word documents that
follow the ARITA (Australian Restructuring Insolvency & Turnaround
Association) standard templates.

The generated documents include review/sign-off placeholders so a Registered
Liquidator can finalise them before lodgement.
"""

from __future__ import annotations

import logging
import os
from datetime import date, datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from models.schemas import (
    AppointmentType,
    CompanyData,
    DIRRIRequest,
    FirmProfile,
)

logger = logging.getLogger("lexsolv.docgen")

# ---------------------------------------------------------------------------
# Output directory
# ---------------------------------------------------------------------------

OUTPUT_DIR = Path(os.getenv("DOCUMENT_OUTPUT_DIR", "generated_documents"))
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# Appointment type display labels
# ---------------------------------------------------------------------------

_APPOINTMENT_LABELS: dict[AppointmentType, str] = {
    AppointmentType.VOLUNTARY_ADMINISTRATION: "Voluntary Administrator",
    AppointmentType.CREDITORS_VOLUNTARY_LIQUIDATION: "Liquidator (Creditors' Voluntary Liquidation)",
    AppointmentType.COURT_LIQUIDATION: "Liquidator (Court Liquidation)",
    AppointmentType.RECEIVERSHIP: "Receiver and Manager",
    AppointmentType.SMALL_BUSINESS_RESTRUCTURING: "Small Business Restructuring Practitioner",
    AppointmentType.DEED_OF_COMPANY_ARRANGEMENT: "Deed Administrator",
}

_APPOINTMENT_ACT_REFS: dict[AppointmentType, str] = {
    AppointmentType.VOLUNTARY_ADMINISTRATION: "Part 5.3A of the Corporations Act 2001",
    AppointmentType.CREDITORS_VOLUNTARY_LIQUIDATION: "Part 5.5 of the Corporations Act 2001",
    AppointmentType.COURT_LIQUIDATION: "Part 5.4 of the Corporations Act 2001",
    AppointmentType.RECEIVERSHIP: "the terms of the relevant security",
    AppointmentType.SMALL_BUSINESS_RESTRUCTURING: "Part 5.3B of the Corporations Act 2001",
    AppointmentType.DEED_OF_COMPANY_ARRANGEMENT: "Part 5.3A of the Corporations Act 2001",
}


# ===================================================================
# Helper: consistent document styling
# ===================================================================

def _set_style_defaults(doc: Document) -> None:
    """Apply firm, professional styling defaults to the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.space_before = Pt(0)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)


def _add_field(doc: Document, label: str, value: str, bold_label: bool = True) -> None:
    """Add a label: value pair as a paragraph."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = bold_label
    run_label.font.size = Pt(11)
    run_value = p.add_run(value)
    run_value.font.size = Pt(11)


def _add_signature_block(doc: Document, name: str, title: str, firm: str) -> None:
    """Add a signature placeholder block for the practitioner."""
    doc.add_paragraph()  # spacer
    p = doc.add_paragraph()
    p.add_run("_" * 50)
    p.paragraph_format.space_after = Pt(2)

    sig_name = doc.add_paragraph()
    run = sig_name.add_run(name)
    run.bold = True
    run.font.size = Pt(11)
    sig_name.paragraph_format.space_after = Pt(0)

    sig_title = doc.add_paragraph()
    run = sig_title.add_run(title)
    run.font.size = Pt(10)
    run.font.italic = True
    sig_title.paragraph_format.space_after = Pt(0)

    sig_firm = doc.add_paragraph()
    run = sig_firm.add_run(firm)
    run.font.size = Pt(10)
    sig_firm.paragraph_format.space_after = Pt(0)

    sig_date = doc.add_paragraph()
    run = sig_date.add_run("Date: ______ / ______ / ________")
    run.font.size = Pt(10)


def _add_review_placeholder(doc: Document, section_name: str) -> None:
    """Add a highlighted review placeholder box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(f"[PRACTITIONER REVIEW REQUIRED — {section_name}]")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)  # red accent


def _format_date(d: date) -> str:
    """Format a date for Australian documents."""
    return d.strftime("%-d %B %Y")


def _company_description(company: CompanyData) -> str:
    """Build a short company identifier string."""
    parts = [company.legal_name]
    if company.acn:
        parts.append(f"(ACN {company.acn[:3]} {company.acn[3:6]} {company.acn[6:]})")
    elif company.abn:
        parts.append(f"(ABN {company.abn})")
    return " ".join(parts)


# ===================================================================
# DocumentGenerator
# ===================================================================

class DocumentGenerator:
    """
    Generates statutory insolvency documents in .docx format.

    All generated files are saved to ``OUTPUT_DIR`` and can be served by the
    FastAPI static-files mount.
    """

    # ---------------------------------------------------------------
    # DIRRI
    # ---------------------------------------------------------------

    def generate_dirri(self, request: DIRRIRequest) -> Path:
        """
        Generate a DIRRI document compliant with Section 506A of the
        Corporations Act 2001 and ARITA's Code of Professional Practice.

        Returns the ``Path`` of the generated .docx file.
        """
        doc = Document()
        _set_style_defaults(doc)

        firm = request.firm_profile
        company = request.company
        appointment_label = _APPOINTMENT_LABELS.get(
            request.appointment_type, "Administrator"
        )
        act_ref = _APPOINTMENT_ACT_REFS.get(
            request.appointment_type, "the Corporations Act 2001"
        )
        company_desc = _company_description(company)

        # ── Document Title ────────────────────────────────────────
        title = doc.add_heading(
            "Declaration of Independence, Relevant Relationships\nand Indemnities (DIRRI)",
            level=0,
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            run.font.size = Pt(20)

        # Sub-title
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(
            f"Pursuant to Section 506A of the Corporations Act 2001\n"
            f"and the Insolvency Practice Rules (Corporations) 2016"
        )
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

        doc.add_paragraph()  # spacer

        # ── Part A: Appointment Details ───────────────────────────
        _add_heading(doc, "Part A — Appointment Details", level=1)

        _add_field(doc, "Company", company_desc)
        if company.trading_name:
            _add_field(doc, "Trading As", company.trading_name)
        _add_field(doc, "Appointment Type", appointment_label)
        _add_field(doc, "Date of Appointment", _format_date(request.appointment_date))
        _add_field(doc, "Appointed By", request.appointing_body)
        _add_field(doc, "Legislative Reference", act_ref)

        doc.add_paragraph()

        # ── Part B: Practitioner Details ──────────────────────────
        _add_heading(doc, "Part B — Practitioner Details", level=1)

        _add_field(doc, "Practitioner", firm.practitioner_name)
        if firm.practitioner_registration_number:
            _add_field(doc, "ASIC Registration No.", firm.practitioner_registration_number)
        _add_field(doc, "Firm", firm.firm_name)
        if firm.firm_abn:
            _add_field(doc, "Firm ABN", firm.firm_abn)
        if firm.firm_address:
            _add_field(doc, "Address", firm.firm_address)
        if firm.firm_phone:
            _add_field(doc, "Phone", firm.firm_phone)
        if firm.firm_email:
            _add_field(doc, "Email", firm.firm_email)

        doc.add_paragraph()

        # ── Part C: Declaration of Independence ───────────────────
        _add_heading(doc, "Part C — Declaration of Independence", level=1)

        doc.add_paragraph(
            f"I, {firm.practitioner_name}, of {firm.firm_name}, have undertaken "
            f"a proper assessment of the risks to my independence prior to accepting "
            f"the appointment as {appointment_label} of {company_desc} in accordance "
            f"with the law and applicable professional standards."
        )

        doc.add_paragraph(
            "This assessment identified no real or potential risks to my acting "
            "independently, other than those set out in this declaration."
            if not request.prior_professional_relationship and not request.relevant_relationships
            else "This assessment identified the following matters which I consider "
                 "relevant to my independence. I have set out the details of each matter "
                 "and the steps I have taken or propose to take to address them."
        )

        doc.add_paragraph(
            "I am not aware of any reasons that would prevent me from accepting "
            "this appointment."
        )

        _add_review_placeholder(doc, "Confirm independence declaration is accurate")

        doc.add_paragraph()

        # ── Part D: Prior Professional Relationships ──────────────
        _add_heading(doc, "Part D — Prior Professional Relationships", level=1)

        if request.prior_professional_relationship and request.prior_relationship_details:
            doc.add_paragraph(
                f"The practitioner has had the following prior professional relationship "
                f"with the Company, its directors, or associated entities:"
            )
            doc.add_paragraph(request.prior_relationship_details)
            doc.add_paragraph(
                "Notwithstanding this prior relationship, I believe that this does not "
                "affect my ability to act independently. My reasons are set out below:"
            )
            _add_review_placeholder(doc, "Insert reasons why prior relationship does not affect independence")
        else:
            doc.add_paragraph(
                f"Neither I, nor my firm {firm.firm_name}, have had any prior "
                f"professional relationship with {company_desc}, its directors, "
                f"or any associated entities within the preceding 24 months."
            )

        doc.add_paragraph()

        # ── Part E: Relevant Relationships and Declarations ───────
        _add_heading(doc, "Part E — Relevant Relationships and Declarations", level=1)

        if request.relevant_relationships:
            doc.add_paragraph(
                "The following relevant relationships and/or dealings are disclosed:"
            )
            for idx, rel in enumerate(request.relevant_relationships, 1):
                p = doc.add_paragraph(style="List Number")
                p.add_run(rel)
            doc.add_paragraph(
                "I have assessed these relationships and am satisfied that they do not "
                "result in a conflict of interest or duty, and do not create a reasonable "
                "apprehension of a lack of independence."
            )
            _add_review_placeholder(doc, "Confirm all relevant relationships are disclosed")
        else:
            doc.add_paragraph(
                "I have no relevant relationships to declare with the Company, "
                "its directors, shareholders, secured creditors, or any associated entities."
            )

        doc.add_paragraph()

        # ── Part F — Indemnities ──────────────────────────────────
        _add_heading(doc, "Part F — Indemnities", level=1)

        if request.indemnities_received and request.indemnity_details:
            doc.add_paragraph(
                "The following indemnity/indemnities have been received or are expected "
                "in connection with this appointment:"
            )
            doc.add_paragraph(request.indemnity_details)
            _add_review_placeholder(doc, "Verify indemnity details and assess impact on independence")
        else:
            doc.add_paragraph(
                "I have not received, nor expect to receive, any indemnity in connection "
                "with this appointment, whether from a director, shareholder, secured "
                "creditor, or any other party."
            )

        doc.add_paragraph()

        # ── Part G — Up-Front Payments ────────────────────────────
        _add_heading(doc, "Part G — Up-Front Payments", level=1)

        if request.upfront_payments_received and request.upfront_payment_details:
            doc.add_paragraph(
                "The following up-front payment(s) have been received prior to "
                "or at the time of appointment:"
            )
            doc.add_paragraph(request.upfront_payment_details)
            _add_review_placeholder(doc, "Confirm up-front payment disclosure is complete")
        else:
            doc.add_paragraph(
                "No up-front payments have been received from any party in connection "
                "with this appointment."
            )

        doc.add_paragraph()

        # ── Part H — Additional Disclosures ───────────────────────
        if request.additional_notes:
            _add_heading(doc, "Part H — Additional Disclosures", level=1)
            doc.add_paragraph(request.additional_notes)
            _add_review_placeholder(doc, "Review additional disclosures")
            doc.add_paragraph()

        # ── Creditor Information ──────────────────────────────────
        _add_heading(doc, "Information for Creditors", level=1)

        doc.add_paragraph(
            "This document has been prepared in accordance with the requirements of "
            "Section 506A of the Corporations Act 2001 and the Insolvency Practice "
            "Rules (Corporations) 2016, Sub-division B."
        )

        doc.add_paragraph(
            "If you have concerns about the independence of the practitioner or any "
            "matter raised in this declaration, you may:"
        )

        for item in [
            "Raise the matter with the practitioner directly;",
            "Contact the Australian Restructuring Insolvency and Turnaround "
            "Association (ARITA) at www.arita.com.au; or",
            "Lodge a complaint with the Australian Securities and Investments "
            "Commission (ASIC) at www.asic.gov.au.",
        ]:
            doc.add_paragraph(item, style="List Bullet")

        doc.add_paragraph()

        # ── Signature Block ───────────────────────────────────────
        _add_heading(doc, "Declaration and Signature", level=1)

        doc.add_paragraph(
            f"I, {firm.practitioner_name}, hereby declare that the information "
            f"contained in this document is true and correct to the best of my "
            f"knowledge and belief."
        )

        _add_signature_block(
            doc,
            name=firm.practitioner_name,
            title=f"Registered Liquidator — {appointment_label}",
            firm=firm.firm_name,
        )

        # ── Footer ────────────────────────────────────────────────
        doc.add_paragraph()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(
            f"Generated by LexSolv AI on {_format_date(date.today())} — "
            f"DRAFT: For review by {firm.practitioner_name} prior to lodgement"
        )
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # ── Save ──────────────────────────────────────────────────
        safe_name = company.legal_name.replace(" ", "_").replace("/", "-")[:40]
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        filename = f"DIRRI_{safe_name}_{timestamp}.docx"
        filepath = OUTPUT_DIR / filename

        doc.save(str(filepath))
        logger.info("DIRRI generated: %s", filepath)
        return filepath

    # ---------------------------------------------------------------
    # Safe Harbour Assessment (placeholder for future prompt)
    # ---------------------------------------------------------------

    def generate_safe_harbour_checklist(
        self,
        firm: FirmProfile,
        company: CompanyData,
        assessment_date: Optional[date] = None,
    ) -> Path:
        """
        Generate a Safe Harbour (Section 588GA) assessment checklist.

        This is a lighter-weight template showing the key conditions a director
        must satisfy to rely on the safe harbour defence.

        Returns the ``Path`` of the generated .docx file.
        """
        doc = Document()
        _set_style_defaults(doc)

        assessment_date = assessment_date or date.today()
        company_desc = _company_description(company)

        # ── Title ─────────────────────────────────────────────────
        title = doc.add_heading("Safe Harbour Assessment Checklist", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            run.font.size = Pt(20)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(
            "Pursuant to Section 588GA of the Corporations Act 2001"
        )
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

        doc.add_paragraph()

        # ── Company / Practitioner Info ───────────────────────────
        _add_heading(doc, "Assessment Details", level=1)
        _add_field(doc, "Company", company_desc)
        _add_field(doc, "Prepared By", firm.practitioner_name)
        _add_field(doc, "Firm", firm.firm_name)
        _add_field(doc, "Assessment Date", _format_date(assessment_date))

        doc.add_paragraph()

        # ── Checklist ─────────────────────────────────────────────
        _add_heading(doc, "Safe Harbour Conditions — s 588GA(1)", level=1)

        doc.add_paragraph(
            "For the safe harbour defence to apply, the following conditions must "
            "be satisfied at the time the director begins to suspect insolvency. "
            "Tick each item and provide supporting evidence."
        )

        checklist_items = [
            (
                "Proper books and records",
                "The company is maintaining financial records that comply with "
                "Section 286 of the Corporations Act 2001.",
            ),
            (
                "Employee entitlements current",
                "All employee entitlements (wages, superannuation, leave) due and "
                "payable are being paid on time.",
            ),
            (
                "Tax reporting obligations",
                "The company is complying with its tax reporting obligations under "
                "taxation law.",
            ),
            (
                "Course of action reasonably likely to lead to a better outcome",
                "The director(s) have started developing, or are implementing, a "
                "course of action that is reasonably likely to lead to a better "
                "outcome for the company than immediate administration or liquidation.",
            ),
            (
                "Qualified advice obtained",
                "The director(s) have obtained advice from an appropriately qualified "
                "entity (e.g. a registered liquidator, accountant, or restructuring "
                "adviser).",
            ),
            (
                "Proper informing of officers",
                "The director(s) are properly informing themselves of the company's "
                "financial position.",
            ),
        ]

        # Build a table for the checklist
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_cells = table.rows[0].cells
        for i, heading_text in enumerate(["#", "Condition", "Status"]):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(1.2)
            row.cells[1].width = Cm(12)
            row.cells[2].width = Cm(3)

        for idx, (title_text, description) in enumerate(checklist_items, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = row.cells[1]
            cell.paragraphs[0].clear()
            run_title = cell.paragraphs[0].add_run(title_text + "\n")
            run_title.bold = True
            run_title.font.size = Pt(10)
            run_desc = cell.paragraphs[0].add_run(description)
            run_desc.font.size = Pt(9)
            run_desc.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

            status_cell = row.cells[2]
            status_p = status_cell.paragraphs[0]
            status_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = status_p.add_run("☐ Met   ☐ Not Met")
            run.font.size = Pt(9)

        doc.add_paragraph()

        # ── Evidence & Notes ──────────────────────────────────────
        _add_heading(doc, "Evidence and Supporting Notes", level=1)

        _add_review_placeholder(doc, "Attach or reference supporting documents for each condition above")

        doc.add_paragraph(
            "1. ____________________________________________________________"
        )
        doc.add_paragraph(
            "2. ____________________________________________________________"
        )
        doc.add_paragraph(
            "3. ____________________________________________________________"
        )

        doc.add_paragraph()

        # ── Recommendation ────────────────────────────────────────
        _add_heading(doc, "Practitioner Recommendation", level=1)

        doc.add_paragraph(
            "Based on the above assessment, the practitioner's preliminary recommendation is:"
        )

        for option in [
            "☐  Safe harbour conditions ARE satisfied — directors may continue trading "
            "while developing a restructuring plan.",
            "☐  Safe harbour conditions are NOT fully satisfied — immediate action "
            "required to address deficiencies (see notes above).",
            "☐  Safe harbour is NOT available — recommend formal insolvency appointment.",
        ]:
            p = doc.add_paragraph(option)
            p.paragraph_format.space_after = Pt(8)

        _add_review_placeholder(doc, "Select recommendation and provide reasoning")

        doc.add_paragraph()

        # ── Signature ─────────────────────────────────────────────
        _add_heading(doc, "Sign-Off", level=1)

        _add_signature_block(
            doc,
            name=firm.practitioner_name,
            title="Registered Liquidator",
            firm=firm.firm_name,
        )

        # ── Footer ────────────────────────────────────────────────
        doc.add_paragraph()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(
            f"Generated by LexSolv AI on {_format_date(date.today())} — "
            f"DRAFT: For review by {firm.practitioner_name} prior to use"
        )
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # ── Save ──────────────────────────────────────────────────
        safe_name = company.legal_name.replace(" ", "_").replace("/", "-")[:40]
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        filename = f"SafeHarbour_{safe_name}_{timestamp}.docx"
        filepath = OUTPUT_DIR / filename

        doc.save(str(filepath))
        logger.info("Safe Harbour checklist generated: %s", filepath)
        return filepath
