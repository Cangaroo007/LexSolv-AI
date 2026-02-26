"""
Universal Document Ingester — extracts raw content from any supported file type.

Accepts PDF, image, DOCX, Excel, and CSV files and returns raw text, tables,
and images for downstream parsing. Does NOT interpret fields — that is the
parser's job (services/file_parser.py).
"""

from __future__ import annotations

import base64
import io
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd


# ---------------------------------------------------------------------------
# Custom exceptions
# ---------------------------------------------------------------------------

class UnsupportedFileTypeError(Exception):
    """Raised when the uploaded file type is not supported."""


class EmptyFileError(Exception):
    """Raised when the uploaded file is zero bytes."""


# ---------------------------------------------------------------------------
# Raw content container
# ---------------------------------------------------------------------------

@dataclass
class RawDocumentContent:
    filename: str
    file_type: str                    # pdf | image | docx | excel | csv
    text_content: str                 # Full extracted text
    tables: list[list[list[str]]]     # List of tables, each a list of rows
    images_base64: list[str]          # For scanned PDFs or image uploads
    metadata: dict                    # Page count, sheet names, doc properties
    likely_scanned: bool              # True if PDF has sparse text
    is_structured: bool               # True if Xero/MYOB columns detected
    raw_bytes: bytes                  # Original file — needed by structured parser


# ---------------------------------------------------------------------------
# Extension → handler mapping
# ---------------------------------------------------------------------------

_PDF_EXTENSIONS = {".pdf"}
_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".heic", ".tiff"}
_DOCX_EXTENSIONS = {".docx"}
_EXCEL_EXTENSIONS = {".xlsx", ".xls"}
_CSV_EXTENSIONS = {".csv"}

SUPPORTED_EXTENSIONS = (
    _PDF_EXTENSIONS | _IMAGE_EXTENSIONS | _DOCX_EXTENSIONS
    | _EXCEL_EXTENSIONS | _CSV_EXTENSIONS
)

# Column patterns that indicate a Xero or MYOB export
_STRUCTURED_COLUMNS = {
    # Xero aged payables / receivables
    "contact", "current", "30 days", "60 days", "90+ days", "total",
    # MYOB
    "co./last name", "balance due",
    # Xero balance sheet / P&L
    "account", "amount",
}

_SPARSE_TEXT_THRESHOLD = 100  # chars per page below which we flag as scanned


# ---------------------------------------------------------------------------
# Ingester
# ---------------------------------------------------------------------------

class DocumentIngester:
    """
    Accepts any file type and extracts raw content for downstream parsing.
    Does NOT interpret fields — returns raw text, tables, and images only.
    Returns: RawDocumentContent with text, tables, images, and file metadata.
    """

    def ingest(
        self,
        file_bytes: bytes,
        filename: str,
        content_type: str = "",
    ) -> RawDocumentContent:
        """Route to correct extractor based on file type."""
        if len(file_bytes) == 0:
            raise EmptyFileError(f"File '{filename}' is empty (0 bytes).")

        ext = Path(filename).suffix.lower()

        if ext in _PDF_EXTENSIONS:
            return self._ingest_pdf(file_bytes, filename)
        if ext in _IMAGE_EXTENSIONS:
            return self._ingest_image(file_bytes, filename)
        if ext in _DOCX_EXTENSIONS:
            return self._ingest_docx(file_bytes, filename)
        if ext in _EXCEL_EXTENSIONS:
            return self._ingest_excel_nonstandard(file_bytes, filename)
        if ext in _CSV_EXTENSIONS:
            return self._ingest_csv(file_bytes, filename)

        raise UnsupportedFileTypeError(
            f"File type '{ext}' is not supported. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def _ingest_pdf(
        self, file_bytes: bytes, filename: str = "document.pdf"
    ) -> RawDocumentContent:
        """
        Use pdfplumber to extract:
        - Full text (page by page)
        - Tables detected by pdfplumber (as list[list[str]])
        - Page count and metadata
        If pdfplumber finds no tables and text is sparse (<100 chars/page),
        flag as likely_scanned=True and include page images as base64.
        """
        import pdfplumber

        all_text_parts: list[str] = []
        all_tables: list[list[list[str]]] = []
        images_b64: list[str] = []
        page_count = 0

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                all_text_parts.append(text)

                for table in page.extract_tables() or []:
                    # Normalise cells to strings
                    cleaned = [
                        [str(cell) if cell is not None else "" for cell in row]
                        for row in table
                    ]
                    all_tables.append(cleaned)

        full_text = "\n".join(all_text_parts)

        # Detect scanned documents
        avg_chars = len(full_text.strip()) / max(page_count, 1)
        likely_scanned = (
            len(all_tables) == 0 and avg_chars < _SPARSE_TEXT_THRESHOLD
        )

        if likely_scanned:
            # Convert each page to an image for Claude vision
            images_b64 = self._pdf_pages_to_base64(file_bytes)

        return RawDocumentContent(
            filename=filename,
            file_type="pdf",
            text_content=full_text,
            tables=all_tables,
            images_base64=images_b64,
            metadata={"page_count": page_count},
            likely_scanned=likely_scanned,
            is_structured=False,
            raw_bytes=file_bytes,
        )

    @staticmethod
    def _pdf_pages_to_base64(file_bytes: bytes) -> list[str]:
        """Render each PDF page as a JPEG image and return base64 strings."""
        import pdfplumber
        from PIL import Image as PILImage

        images: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                img = page.to_image(resolution=200).original
                # Resize to Claude optimal max dimension
                img = _resize_to_max(img, 1568)
                buf = io.BytesIO()
                img.save(buf, format="JPEG", quality=85)
                images.append(base64.b64encode(buf.getvalue()).decode())
        return images

    # ------------------------------------------------------------------
    # Image
    # ------------------------------------------------------------------

    def _ingest_image(
        self, file_bytes: bytes, filename: str
    ) -> RawDocumentContent:
        """
        Accept: jpg, jpeg, png, webp, heic, tiff.
        Convert to JPEG via Pillow if needed (Claude accepts jpeg/png/gif/webp).
        Return as base64-encoded image ready for Claude vision.
        Resize to max 1568px on longest side (Claude optimal size).
        """
        from PIL import Image as PILImage

        img = PILImage.open(io.BytesIO(file_bytes))
        img = _resize_to_max(img, 1568)

        # Convert to RGB if necessary (e.g. RGBA PNGs, palette images)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        b64 = base64.b64encode(buf.getvalue()).decode()

        return RawDocumentContent(
            filename=filename,
            file_type="image",
            text_content="",
            tables=[],
            images_base64=[b64],
            metadata={
                "original_size": list(PILImage.open(io.BytesIO(file_bytes)).size),
                "resized_size": list(img.size),
            },
            likely_scanned=False,
            is_structured=False,
            raw_bytes=file_bytes,
        )

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _ingest_docx(
        self, file_bytes: bytes, filename: str = "document.docx"
    ) -> RawDocumentContent:
        """
        Use python-docx to extract:
        - All paragraph text (preserving heading levels)
        - All tables (as list[list[str]])
        - Document properties (title, author, created date)
        """
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))

        # Extract paragraphs with heading level info
        text_parts: list[str] = []
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                text_parts.append(f"[H{level}] {para.text}")
            else:
                text_parts.append(para.text)

        # Extract tables
        all_tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            all_tables.append(rows)

        # Document properties
        props = doc.core_properties
        metadata = {
            "title": props.title or "",
            "author": props.author or "",
            "created": str(props.created) if props.created else "",
        }

        return RawDocumentContent(
            filename=filename,
            file_type="docx",
            text_content="\n".join(text_parts),
            tables=all_tables,
            images_base64=[],
            metadata=metadata,
            likely_scanned=False,
            is_structured=False,
            raw_bytes=file_bytes,
        )

    # ------------------------------------------------------------------
    # Excel (non-standard)
    # ------------------------------------------------------------------

    def _ingest_excel_nonstandard(
        self, file_bytes: bytes, filename: str = "spreadsheet.xlsx"
    ) -> RawDocumentContent:
        """
        Use pandas to read ALL sheets (not just sheet 0).
        Return each sheet as a table with sheet name as label.
        Mark as nonstandard=True to signal AI parser is needed.
        """
        sheets = pd.read_excel(
            io.BytesIO(file_bytes), sheet_name=None, dtype=str
        )

        all_tables: list[list[list[str]]] = []
        sheet_names: list[str] = []

        for name, df in sheets.items():
            df = df.fillna("")
            header_row = [str(c) for c in df.columns.tolist()]
            data_rows = df.astype(str).values.tolist()
            table = [header_row] + data_rows
            all_tables.append(table)
            sheet_names.append(str(name))

        # Combine all text for quick preview
        text_content = "\n".join(
            f"--- Sheet: {name} ---\n"
            + "\n".join("\t".join(row) for row in tbl)
            for name, tbl in zip(sheet_names, all_tables)
        )

        return RawDocumentContent(
            filename=filename,
            file_type="excel",
            text_content=text_content,
            tables=all_tables,
            images_base64=[],
            metadata={"sheet_names": sheet_names, "sheet_count": len(sheet_names)},
            likely_scanned=False,
            is_structured=False,
            raw_bytes=file_bytes,
        )

    # ------------------------------------------------------------------
    # CSV
    # ------------------------------------------------------------------

    def _ingest_csv(
        self, file_bytes: bytes, filename: str = "data.csv"
    ) -> RawDocumentContent:
        """
        Read with pandas. If column headers match known Xero/MYOB patterns,
        mark as is_structured=True (structured parser will handle it).
        Otherwise mark as nonstandard=True.
        """
        df = pd.read_csv(io.BytesIO(file_bytes), dtype=str).fillna("")

        header_row = [str(c) for c in df.columns.tolist()]
        data_rows = df.astype(str).values.tolist()
        table = [header_row] + data_rows

        text_content = "\n".join("\t".join(row) for row in table)

        # Check for structured accounting software columns
        lower_cols = {c.strip().lower() for c in header_row}
        is_structured = bool(lower_cols & _STRUCTURED_COLUMNS)

        return RawDocumentContent(
            filename=filename,
            file_type="csv",
            text_content=text_content,
            tables=[table],
            images_base64=[],
            metadata={"columns": header_row, "row_count": len(data_rows)},
            likely_scanned=False,
            is_structured=is_structured,
            raw_bytes=file_bytes,
        )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _resize_to_max(img, max_px: int = 1568):
    """Resize image so longest side is at most *max_px*, preserving aspect ratio."""
    from PIL import Image as PILImage

    w, h = img.size
    if max(w, h) <= max_px:
        return img
    scale = max_px / max(w, h)
    new_size = (int(w * scale), int(h * scale))
    return img.resize(new_size, PILImage.LANCZOS)
